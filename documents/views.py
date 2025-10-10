from django.shortcuts import render
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.http import FileResponse, Http404, JsonResponse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
import os
import json
import uuid
from datetime import datetime
import pytesseract
import cv2
import numpy as np
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from ollama_integration import AIDocumentProcessor, OllamaClient
from PIL import Image
import fitz  # PyMuPDF for PDF processing
# Use enhanced field detector for better field recognition
try:
    from enhanced_field_detector import EnhancedFieldDetector, convert_form_fields_to_dict
    EnhancedDocumentProcessor = EnhancedFieldDetector
    print("Using Enhanced Field Detector")
except ImportError:
    try:
        from improved_field_detector import ImprovedFieldDetector, convert_form_fields_to_dict
        EnhancedDocumentProcessor = ImprovedFieldDetector
        print("Using Improved Field Detector")
    except ImportError:
        try:
            from enhanced_document_processor import EnhancedDocumentProcessor, convert_form_fields_to_dict
            print("Using Enhanced Document Processor")
        except ImportError:
            from simple_enhanced_processor import SimpleEnhancedProcessor as EnhancedDocumentProcessor, convert_form_fields_to_dict
            print("Using Simple Enhanced Processor")
from intelligent_field_filler import IntelligentFieldFiller

# Import Universal Document Processor for automatic training
try:
    from universal_document_processor import UniversalDocumentProcessor
    universal_processor = UniversalDocumentProcessor()
    print("Universal Document Processor loaded for automatic training")
except ImportError:
    universal_processor = None
    print("Universal Document Processor not available")

# Configure Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Import persistent training storage
from training_storage import training_storage

# In-memory storage (no database needed)
documents_storage = {}
chat_sessions = {}

def auto_train_from_document(file_path, document, result):
    """
    Automatically train the Universal Document Processor from uploaded document
    """
    if not universal_processor:
        print("Universal processor not available")
        return None
    
    try:
        print(f"Starting auto training for: {file_path}")
        
        # Extract text from document
        text = result.get('extracted_text', '')
        if not text:
            print("No extracted text found")
            return None
        
        print(f"Text length: {len(text)}")
        
        # Classify document type
        doc_type, confidence = universal_processor.classify_document_type(text)
        print(f"Auto-classified document type: {doc_type} (confidence: {confidence:.2f})")
        
        # Create training data from detected fields
        training_samples = []
        fields = document.get('fields', [])
        print(f"Processing {len(fields)} fields")
        print(f"Document keys: {list(document.keys())}")
        print(f"First few fields: {fields[:3] if fields else 'No fields'}")
        
        for i, field in enumerate(fields):
            context = field.get('context', '')
            confidence = field.get('confidence', 0)
            field_type = field.get('field_type', 'text')
            field_id = field.get('id', '')
            
            # Handle both numeric ID and string field_id
            if isinstance(field_id, int):
                field_id = str(field_id)
            
            # For stored documents, the context field contains the technical field ID
            # Extract it for field type mapping
            technical_field_id = context if context else ''
            
            print(f"Field {i}: context='{context}' confidence={confidence} type={field_type} id='{field_id}'")
            
            # For PDF form fields, use the field ID as context if no other context
            if not context or context.strip() == '':
                if field_id:
                    # For PDF form fields, extract meaningful field name from field ID
                    if field_id.startswith('topmostSubform'):
                        # Extract field identifier from technical ID (e.g., f1_01 from topmostSubform[0].Page1[0].Step1a[0].f1_01[0])
                        field_parts = field_id.split('.')
                        if field_parts:
                            last_part = field_parts[-1]
                            if last_part.startswith('f') and '_' in last_part:
                                context = last_part  # Use f1_01, f1_02, etc. as context
                            else:
                                context = f"Field {i+1}"
                        else:
                            context = f"Field {i+1}"
                    else:
                        context = field_id
                else:
                    # Generate context from field ID for form fields
                    context = f"Field {i+1}"
            
            # Always create training samples for detected fields
            # Map field types to universal types using the technical field ID
            universal_field_type = map_field_type_to_universal(field_type, technical_field_id)
            
            # Create training sample with minimum confidence of 0.5
            training_sample = {
                'text': context,
                'field_type': universal_field_type,
                'document_type': doc_type,
                'context': context,
                'confidence': max(0.5, confidence)  # Ensure minimum confidence
            }
            training_samples.append(training_sample)
            print(f"Added training sample: {training_sample}")
        
        if not training_samples:
            print("No valid training samples found")
            return None
        
        print(f"Created {len(training_samples)} training samples")
        
        # Always add samples to training data
        for sample in training_samples:
            universal_processor.training_data.append(sample)
        
        # Also save to persistent storage
        training_storage.add_training_samples(training_samples)
        
        print(f"Added {len(training_samples)} samples to training data. Total samples: {len(universal_processor.training_data)}")
        
        # Train the model (only if we have enough samples)
        if len(universal_processor.training_data) >= 3:
            print("Training model with all samples...")
            training_results = universal_processor.train_model(universal_processor.training_data)
            print(f"Training results: {training_results}")
            
            return {
                'samples_added': len(training_samples),
                'document_type': doc_type,
                'field_type_accuracy': training_results.get('field_type_accuracy', 0),
                'document_type_accuracy': training_results.get('document_type_accuracy', 0),
                'training_samples': training_results.get('training_samples', len(universal_processor.training_data))
            }
        else:
            return {
                'samples_added': len(training_samples),
                'document_type': doc_type,
                'training_samples': len(universal_processor.training_data),
                'note': f'Added to training data. Total samples: {len(universal_processor.training_data)} (need 3+ for retraining)'
            }
    
    except Exception as e:
        print(f"Error in automatic training: {e}")
        import traceback
        traceback.print_exc()
        return None

def map_field_type_to_universal(field_type, field_id=''):
    """
    Map detected field types to universal field types
    """
    # Ensure field_id is a string
    if isinstance(field_id, int):
        field_id = str(field_id)
    elif not isinstance(field_id, str):
        field_id = str(field_id) if field_id else ''
    mapping = {
        'name': 'name',
        'email': 'email', 
        'phone': 'phone',
        'address': 'address',
        'date': 'date_of_birth',
        'ssn': 'ssn',
        'id_number': 'id_number',
        'income': 'income',
        'bank_account': 'bank_account',
        'credit_score': 'credit_score',
        'assets': 'assets',
        'patient_id': 'patient_id',
        'insurance': 'insurance',
        'medications': 'medications',
        'allergies': 'allergies',
        'case_number': 'case_number',
        'court': 'court',
        'attorney': 'attorney',
        'student_id': 'student_id',
        'gpa': 'gpa',
        'major': 'major',
        'company_name': 'company_name',
        'position': 'position',
        'experience_years': 'experience_years',
        'salary': 'expected_salary'
    }
    
    # Try direct mapping first
    if field_type in mapping:
        return mapping[field_type]
    
    # For PDF form fields, try to infer from field ID
    if field_id:
        field_id_lower = field_id.lower()
        
        # Extract the field identifier part (e.g., f1_01 from topmostSubform[0].Page1[0].Step1a[0].f1_01[0])
        field_identifier = field_id
        if field_id.startswith('topmostSubform'):
            field_parts = field_id.split('.')
            if field_parts:
                field_identifier = field_parts[-1]
        
        # W-4 form specific mappings using field identifiers
        if 'name' in field_id_lower or 'f1_01' in field_identifier:
            return 'name'
        elif 'ssn' in field_id_lower or 'f1_02' in field_identifier:
            return 'ssn'
        elif 'address' in field_id_lower or 'f1_03' in field_identifier:
            return 'address'
        elif 'city' in field_id_lower or 'f1_04' in field_identifier:
            return 'address'
        elif 'state' in field_id_lower or 'f1_05' in field_identifier:
            return 'address'
        elif 'zip' in field_id_lower or 'f1_06' in field_identifier:
            return 'address'
        elif 'phone' in field_id_lower:
            return 'phone'
        elif 'email' in field_id_lower:
            return 'email'
        elif 'date' in field_id_lower:
            return 'date_of_birth'
        elif 'income' in field_id_lower or 'wage' in field_id_lower:
            return 'income'
        elif 'dependent' in field_id_lower:
            return 'number'
        elif 'allowance' in field_id_lower:
            return 'number'
    
    return 'text'

class DocumentProcessor:
    def __init__(self):
        self.blank_spaces = []
        self.extracted_text = ""
        self._enhanced_processor = None
        self._intelligent_filler = None
        self.ai_processor = AIDocumentProcessor()
    
    @property
    def enhanced_processor(self):
        if self._enhanced_processor is None:
            self._enhanced_processor = EnhancedDocumentProcessor()
        return self._enhanced_processor
    
    @property
    def intelligent_filler(self):
        if self._intelligent_filler is None:
            self._intelligent_filler = IntelligentFieldFiller()
        return self._intelligent_filler
        
    def pdf_to_images(self, pdf_path):
        """Convert all PDF pages to images for processing"""
        try:
            # Open PDF
            pdf_document = fitz.open(pdf_path)
            images = []
            
            # Convert each page to image
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Convert to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to OpenCV format
                nparr = np.frombuffer(img_data, np.uint8)
                image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                images.append((page_num, image))
            
            pdf_document.close()
            return images
        except Exception as e:
            raise Exception(f"Error converting PDF to images: {str(e)}")
    
    def process_document(self, file_path):
        """Process uploaded document using enhanced field detection"""
        try:
            # Use enhanced processor for better field detection
            result = self.enhanced_processor.process_document(file_path)
            
            # Convert FormField objects to dictionary format
            self.blank_spaces = convert_form_fields_to_dict(result['fields'])
            self.extracted_text = result['extracted_text']
            
            return {
                'extracted_text': self.extracted_text,
                'blank_spaces': self.blank_spaces,
                'total_blanks': len(self.blank_spaces)
            }
        except Exception as e:
            # Fallback to original method if enhanced processor fails
            print(f"Enhanced processor failed, falling back to original method: {e}")
            import traceback
            traceback.print_exc()
            return self._process_document_fallback(file_path)
    
    def _process_document_fallback(self, file_path):
        """Fallback document processing method - handles multi-page PDFs"""
        try:
            # Check file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            
            all_blank_spaces = []
            all_extracted_text = []
            
            if file_ext == '.pdf':
                # Process all PDF pages
                images = self.pdf_to_images(file_path)
                
                for page_num, image in images:
                    # Convert to grayscale
                    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                    
                    # Apply OCR to extract text
                    page_text = pytesseract.image_to_string(gray)
                    all_extracted_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    # Find blank spaces (white rectangles)
                    blank_spaces = self.find_blank_spaces(gray, page_num)
                    
                    # If no blank spaces found, create virtual fields based on text analysis
                    if not blank_spaces:
                        blank_spaces = self.create_virtual_fields_from_text(page_text, gray, page_num)
                    
                    all_blank_spaces.extend(blank_spaces)
                
                self.extracted_text = '\n'.join(all_extracted_text)
                self.blank_spaces = all_blank_spaces
            else:
                # Process image file (single page)
                image = cv2.imread(file_path)
                
                # Check if image was loaded successfully
                if image is None:
                    raise ValueError(f"Could not load image from {file_path}. Please ensure the file is a valid image or PDF format.")
                
                # Convert to grayscale
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                
                # Apply OCR to extract text
                self.extracted_text = pytesseract.image_to_string(gray)
                
                # Find blank spaces (white rectangles)
                self.blank_spaces = self.find_blank_spaces(gray, 0)
                
                # If no blank spaces found, create virtual fields based on text analysis
                if not self.blank_spaces:
                    self.blank_spaces = self.create_virtual_fields_from_text(self.extracted_text, gray, 0)
            
            return {
                'extracted_text': self.extracted_text,
                'blank_spaces': self.blank_spaces,
                'total_blanks': len(self.blank_spaces)
            }
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def find_blank_spaces(self, gray_image, page_num=0):
        """Find blank spaces in the document"""
        try:
            # Apply adaptive threshold for better edge detection
            thresh = cv2.adaptiveThreshold(gray_image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2)
            
            # Find contours
            contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            blank_spaces = []
            image_height, image_width = gray_image.shape
            
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                area = w * h
                
                # More specific filtering for form fields
                # Look for rectangular areas that could be form fields
                aspect_ratio = w / h if h > 0 else 0
                
                # Filter for reasonable form field sizes and shapes
                if (area > 1000 and area < 100000 and  # Larger minimum size to avoid artifacts
                    w > 50 and h > 20 and  # Larger minimum dimensions
                    w < image_width * 0.8 and h < image_height * 0.3 and  # Not too large
                    (aspect_ratio > 0.5 and aspect_ratio < 10)):  # Reasonable aspect ratio
                    
                    # Check if this area is actually blank (mostly white)
                    roi = gray_image[y:y+h, x:x+w]
                    if roi.size > 0:
                        mean_intensity = np.mean(roi)
                        if mean_intensity > 200:  # Mostly white/blank area
                            blank_spaces.append({
                                'x': int(x),
                                'y': int(y),
                                'width': int(w),
                                'height': int(h),
                                'area': int(area),
                                'context': analyze_context(gray_image, x, y, w, h, self.extracted_text),
                                'page': page_num
                            })
            
            # If no blank spaces found with the above method, try a simpler approach
            if not blank_spaces:
                # Look for white rectangular regions
                _, thresh_white = cv2.threshold(gray_image, 240, 255, cv2.THRESH_BINARY)
                contours_white, _ = cv2.findContours(thresh_white, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                for contour in contours_white:
                    x, y, w, h = cv2.boundingRect(contour)
                    area = w * h
                    
                    if (area > 1000 and area < 50000 and
                        w > 50 and h > 20 and
                        w < image_width * 0.6 and h < image_height * 0.2):
                        
                        blank_spaces.append({
                            'x': int(x),
                            'y': int(y),
                            'width': int(w),
                            'height': int(h),
                            'area': int(area),
                            'context': analyze_context(gray_image, x, y, w, h, self.extracted_text),
                            'page': page_num
                        })
            
            return blank_spaces
            
        except Exception as e:
            print(f"Error in find_blank_spaces: {e}")
            return []
    
    def create_virtual_fields_from_text(self, text, gray_image, page_num=0):
        """Create virtual form fields based on text analysis"""
        virtual_fields = []
        lines = text.split('\n')
        
        # Enhanced form field patterns with more specific matching
        field_patterns = {
            'name': ['enter your name', 'please enter your name', 'name of dependent', 'full name'],
            'age': ['age of dependent', 'age:', 'years old'],
            'dropdown': ['select an item', 'dropdown', 'combo', 'choose'],
            'checkbox': ['check all that apply', 'option 1', 'option 2', 'option 3'],
            'email': ['email', 'e-mail', 'email address'],
            'phone': ['phone', 'telephone', 'tel', 'mobile'],
            'address': ['address', 'street', 'location'],
            'date': ['date', 'birth', 'dob'],
            'signature': ['signature', 'sign', 'initial']
        }
        
        # Get image dimensions if available
        if gray_image is not None:
            height, width = gray_image.shape
        else:
            height, width = 800, 600  # Default dimensions
        
        # Create virtual fields based on text content
        field_id = 0
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Skip empty lines
            if not line_lower:
                continue
                
            # Check if line contains form field indicators
            field_type = None
            for ftype, patterns in field_patterns.items():
                for pattern in patterns:
                    if pattern in line_lower:
                        field_type = ftype
                        break
                if field_type:
                    break
            
            # If we found a field type, create a virtual field
            if field_type:
                virtual_fields.append({
                    'x': 50,  # Default position
                    'y': 100 + (field_id * 60),  # Spread vertically with more space
                    'width': 400,  # Wider default width
                    'height': 40,   # Taller default height
                    'area': 16000,   # Larger default area
                    'context': field_type,
                    'page': page_num
                })
                field_id += 1
        
        # If still no fields found, create some default ones based on common form elements
        if not virtual_fields:
            # Look for lines that end with colons (common in forms)
            for i, line in enumerate(lines):
                if line.strip().endswith(':') and len(line.strip()) > 3:
                    field_type = 'general'
                    line_lower = line.lower().strip()
                    
                    # Determine field type
                    for ftype, patterns in field_patterns.items():
                        for pattern in patterns:
                            if pattern in line_lower:
                                field_type = ftype
                                break
                    
                    virtual_fields.append({
                        'x': 50,
                        'y': 100 + (len(virtual_fields) * 60),
                        'width': 400,
                        'height': 40,
                        'area': 16000,
                        'context': field_type,
                        'page': page_num
                    })
        
        return virtual_fields
    
    def process_word_document(self, file_path):
        """Process Word documents (.doc, .docx)"""
        try:
            from docx import Document
            
            # Extract text from Word document
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            # Create virtual fields based on text content
            virtual_fields = self.create_virtual_fields_from_text(content, None)
            
            return {
                'extracted_text': content,
                'blank_spaces': virtual_fields,
                'total_blanks': len(virtual_fields)
            }
        except Exception as e:
            # Fallback to simple text extraction if python-docx fails
            try:
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
                
                virtual_fields = self.create_virtual_fields_from_text(content, None)
                
                return {
                    'extracted_text': content,
                    'blank_spaces': virtual_fields,
                    'total_blanks': len(virtual_fields)
                }
            except Exception as e2:
                raise Exception(f"Error processing Word document: {str(e2)}")
    
    def process_text_document(self, file_path):
        """Process text documents (.txt, .rtf)"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Create virtual fields based on text content
            virtual_fields = self.create_virtual_fields_from_text(content, None)
            
            return {
                'extracted_text': content,
                'blank_spaces': virtual_fields,
                'total_blanks': len(virtual_fields)
            }
        except Exception as e:
            raise Exception(f"Error processing text document: {str(e)}")

def index(request):
    """Main page view"""
    document = None
    ollama_status = {'running': False, 'models': []}
    
    # Check if there's a document in session
    if 'current_document_id' in request.session:
        doc_id = request.session['current_document_id']
        document = documents_storage.get(doc_id)
        if not document:
            del request.session['current_document_id']
    
    # Check Ollama status
    try:
        ollama = OllamaClient()
        ollama_status = {
            'running': ollama.is_ollama_running(),
            'models': ollama.list_models() if ollama.is_ollama_running() else []
        }
    except:
        pass
    
    context = {
        'document': document,
        'ollama_status': ollama_status,
        'now': datetime.now()
    }
    
    return render(request, 'index.html', context)

def analyze_context(image, x, y, w, h, full_text=""):
    """Analyze context around blank space to suggest content"""
    # If we have the full extracted text, use it for better analysis
    if full_text:
        lines = full_text.split('\n')
        
        # Look for form field indicators in the text
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check for specific field types
            if any(keyword in line_lower for keyword in ['enter your name', 'please enter your name', 'name:']):
                return 'name'
            elif any(keyword in line_lower for keyword in ['dependent', 'name of dependent']):
                return 'name'
            elif any(keyword in line_lower for keyword in ['age', 'age of dependent']):
                return 'age'
            elif any(keyword in line_lower for keyword in ['select', 'dropdown', 'combo']):
                return 'dropdown'
            elif any(keyword in line_lower for keyword in ['check', 'option']):
                return 'checkbox'
            elif any(keyword in line_lower for keyword in ['email', 'e-mail']):
                return 'email'
            elif any(keyword in line_lower for keyword in ['phone', 'telephone', 'tel']):
                return 'phone'
            elif any(keyword in line_lower for keyword in ['address']):
                return 'address'
            elif any(keyword in line_lower for keyword in ['date', 'birth']):
                return 'date'
            elif any(keyword in line_lower for keyword in ['signature', 'sign']):
                return 'signature'
    
    # Fallback to OCR on context region
    try:
        padding = 50
        y1 = max(0, y - padding)
        y2 = min(image.shape[0], y + h + padding)
        x1 = max(0, x - padding)
        x2 = min(image.shape[1], x + w + padding)
        
        context_region = image[y1:y2, x1:x2]
        context_text = pytesseract.image_to_string(context_region)
        context_lower = context_text.lower()
        
        if 'name' in context_lower:
            return 'name'
        elif 'address' in context_lower:
            return 'address'
        elif 'phone' in context_lower or 'tel' in context_lower:
            return 'phone'
        elif 'email' in context_lower:
            return 'email'
        elif 'date' in context_lower:
            return 'date'
        elif 'signature' in context_lower:
            return 'signature'
    except:
        pass
    
    return 'general'

@api_view(['POST'])
def upload_document(request):
    """Handle document upload and processing"""
    try:
        print("Starting upload_document function")
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        if not file.name:
            return Response({'error': 'No file selected'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check file size (limit to 50MB)
        if file.size > 50 * 1024 * 1024:
            return Response({'error': 'File too large. Maximum size is 50MB.'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get file extension
        file_ext = os.path.splitext(file.name)[1].lower()
        
        # Supported file types
        supported_extensions = {
            '.pdf': 'PDF Document',
            '.png': 'PNG Image',
            '.jpg': 'JPEG Image', 
            '.jpeg': 'JPEG Image',
            '.gif': 'GIF Image',
            '.bmp': 'BMP Image',
            '.tiff': 'TIFF Image',
            '.tif': 'TIFF Image',
            '.webp': 'WebP Image',
            '.doc': 'Word Document',
            '.docx': 'Word Document',
            '.txt': 'Text Document',
            '.rtf': 'Rich Text Document'
        }
        
        if file_ext not in supported_extensions:
            return Response({
                'error': f'Unsupported file type: {file_ext}. Supported types: {", ".join(supported_extensions.keys())}'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Save uploaded file
        filename = f"{doc_id}_{file.name}"
        upload_path = os.path.join(settings.MEDIA_ROOT, 'uploads')
        os.makedirs(upload_path, exist_ok=True)
        filepath = os.path.join(upload_path, filename)
        
        with open(filepath, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        
        # Process document based on file type
        processor = DocumentProcessor()
        
        # Process document using enhanced processor
        result = processor.process_document(filepath)
        
        # Create document record in memory
        document = {
            'id': doc_id,
            'filename': file.name,
            'file_path': filepath,
            'extracted_text': result['extracted_text'],
            'total_blanks': result['total_blanks'],
            'uploaded_at': datetime.now().isoformat(),
            'status': 'processed',
            'fields': []
        }
        
        # Store document ID in session (if available)
        if hasattr(request, 'session'):
            request.session['current_document_id'] = str(doc_id)
        
        # Create field records in memory
        for i, space in enumerate(result['blank_spaces']):
            field = {
                'id': i,
                'field_type': space.get('field_type', space.get('context', 'text')),
                'x_position': space.get('x_position', space.get('x', 0)),
                'y_position': space.get('y_position', space.get('y', 0)),
                'width': space.get('width', 100),
                'height': space.get('height', 25),
                'area': space.get('area', space.get('width', 100) * space.get('height', 25)),
                'context': space.get('context', 'text'),
                'page': space.get('page', 0),  # Include page number
                'user_content': '',
                'ai_suggestion': '',
                'ai_enhanced': False
            }
            document['fields'].append(field)
        
        # Store in memory
        documents_storage[str(doc_id)] = document
        
        # Also save to persistent storage
        training_storage.save_document(str(doc_id), document)
        
        # DISABLED: Automatic training (user should delete unwanted fields first)
        # User will click "Train System" button after cleaning up fields
        # training_results = None
        # if universal_processor:
        #     try:
        #         training_results = auto_train_from_document(file.name, document, result)
        #         print(f"Automatic training completed: {training_results}")
        #     except Exception as training_error:
        #         print(f"Automatic training failed: {training_error}")
        #         import traceback
        #         traceback.print_exc()
        
        response_data = {
            'success': True,
            'document_id': str(doc_id),
            'document': document,
            'result': {
                'extracted_text': result['extracted_text'],
                'total_blanks': result['total_blanks']
            }
        }
        
        # No automatic training - user will click "Train System" after cleanup
        
        return Response(response_data)
    
    except Exception as e:
        print(f"Upload error: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e), 'traceback': traceback.format_exc()}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def get_document(request, doc_id):
    """Get document information"""
    try:
        document = documents_storage.get(str(doc_id))
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        return Response(document)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def update_field(request, doc_id):
    """Update a specific field in the document"""
    try:
        field_id = request.data.get('field_id')
        content = request.data.get('content', '')
        
        document = documents_storage.get(str(doc_id))
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Find and update the field
        for field in document['fields']:
            if field['id'] == field_id:
                field['user_content'] = content
                break
        else:
            return Response({'error': 'Field not found'}, status=status.HTTP_404_NOT_FOUND)
        
        return Response({'success': True})
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def delete_field(request, doc_id):
    """Delete a specific field from the document"""
    try:
        field_id = request.data.get('field_id')
        
        document = documents_storage.get(str(doc_id))
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Convert field_id to both int and string for flexible matching
        try:
            field_id_int = int(field_id)
            field_id_str = str(field_id)
        except (ValueError, TypeError):
            field_id_int = None
            field_id_str = str(field_id)
        
        # Find and remove the field (match by int OR string)
        original_count = len(document['fields'])
        document['fields'] = [
            field for field in document['fields'] 
            if not (field['id'] == field_id or 
                   field['id'] == field_id_int or 
                   field['id'] == field_id_str or
                   str(field['id']) == field_id_str)
        ]
        
        if len(document['fields']) == original_count:
            return Response({'error': 'Field not found', 'searched_id': field_id, 'total_fields': original_count}, status=status.HTTP_404_NOT_FOUND)
        
        # Update persistent storage
        training_storage.save_document(str(doc_id), document)
        
        return Response({
            'success': True, 
            'remaining_fields': len(document['fields']),
            'deleted_id': field_id
        })
        
    except Exception as e:
        import traceback
        return Response({'error': str(e), 'traceback': traceback.format_exc()}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def generate_pdf(request, doc_id):
    """Generate final PDF with filled content"""
    try:
        document = documents_storage.get(str(doc_id))
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Create PDF
        pdf_filename = f"filled_{doc_id}.pdf"
        processed_path = os.path.join(settings.MEDIA_ROOT, 'processed')
        os.makedirs(processed_path, exist_ok=True)
        pdf_path = os.path.join(processed_path, pdf_filename)
        
        # Create PDF with filled information
        c = canvas.Canvas(pdf_path, pagesize=letter)
        width, height = letter
        
        # Title
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, height - 100, f"Filled Document: {document['filename']}")
        
        # Document information
        c.setFont("Helvetica", 12)
        y_position = height - 150
        
        c.drawString(100, y_position, f"Document ID: {doc_id}")
        y_position -= 30
        c.drawString(100, y_position, f"Uploaded: {document['uploaded_at']}")
        y_position -= 30
        c.drawString(100, y_position, f"Total Fields: {document['total_blanks']}")
        y_position -= 50
        
        # Filled fields
        c.setFont("Helvetica-Bold", 14)
        c.drawString(100, y_position, "Filled Fields:")
        y_position -= 30
        
        c.setFont("Helvetica", 12)
        for field in document['fields']:
            if field['user_content']:
                c.drawString(120, y_position, f"Field {field['id']} ({field['context']}): {field['user_content']}")
                y_position -= 25
        
        c.save()
        
        return Response({
            'success': True,
            'pdf_path': pdf_path,
            'download_url': f'/api/documents/download/{doc_id}'
        })
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def download_pdf(request, doc_id):
    """Download the generated PDF"""
    try:
        pdf_filename = f"filled_{doc_id}.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, 'processed', pdf_filename)
        
        if not os.path.exists(pdf_path):
            return Response({'error': 'PDF not found'}, status=status.HTTP_404_NOT_FOUND)
        
        return FileResponse(open(pdf_path, 'rb'), as_attachment=True, filename=pdf_filename)
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def clear_session(request):
    """Clear the current document from session"""
    try:
        # Clear the current document ID from session
        if 'current_document_id' in request.session:
            doc_id = request.session['current_document_id']
            # Optionally remove the document from memory storage
            if str(doc_id) in documents_storage:
                del documents_storage[str(doc_id)]
            del request.session['current_document_id']
        
        return Response({'success': True, 'message': 'Session cleared successfully'})
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def manual_train(request, doc_id):
    """Manually trigger training from a document"""
    try:
        print(f"Manual training request for document: {doc_id}")
        
        if not universal_processor:
            print("Universal Document Processor not available")
            return Response({'error': 'Universal Document Processor not available'}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Try memory first, then persistent storage
        document = documents_storage.get(str(doc_id))
        
        if not document:
            print(f"Document not found in memory, checking persistent storage: {doc_id}")
            document = training_storage.load_document(str(doc_id))
        
        if not document:
            print(f"Document not found in persistent storage either: {doc_id}")
            # Create training data from existing field patterns
            print("Creating training data from existing field patterns...")
            
            # Add some generic training samples based on common field patterns
            generic_samples = [
                {
                    'text': 'Name',
                    'field_type': 'name',
                    'document_type': 'form',
                    'context': 'Name',
                    'confidence': 0.8
                },
                {
                    'text': 'Email Address',
                    'field_type': 'email',
                    'document_type': 'form',
                    'context': 'Email Address',
                    'confidence': 0.8
                },
                {
                    'text': 'Phone Number',
                    'field_type': 'phone',
                    'document_type': 'form',
                    'context': 'Phone Number',
                    'confidence': 0.8
                },
                {
                    'text': 'Social Security Number',
                    'field_type': 'ssn',
                    'document_type': 'form',
                    'context': 'Social Security Number',
                    'confidence': 0.8
                },
                {
                    'text': 'Address',
                    'field_type': 'address',
                    'document_type': 'form',
                    'context': 'Address',
                    'confidence': 0.8
                }
            ]
            
            # Add samples to training data
            for sample in generic_samples:
                universal_processor.training_data.append(sample)
            
            # Also save to persistent storage
            training_storage.add_training_samples(generic_samples)
            
            print(f"Added {len(generic_samples)} generic training samples")
            
            # Train if we have enough samples
            if len(universal_processor.training_data) >= 3:
                training_results = universal_processor.train_model(universal_processor.training_data)
                return Response({
                    'success': True,
                    'message': 'Training completed with generic samples',
                    'training': {
                        'samples_added': len(generic_samples),
                        'document_type': 'form',
                        'training_samples': len(universal_processor.training_data),
                        'field_type_accuracy': training_results.get('field_type_accuracy', 0),
                        'document_type_accuracy': training_results.get('document_type_accuracy', 0)
                    }
                })
            else:
                return Response({
                    'success': True,
                    'message': 'Added generic training samples',
                    'training': {
                        'samples_added': len(generic_samples),
                        'document_type': 'form',
                        'training_samples': len(universal_processor.training_data),
                        'note': 'Need more samples for model training'
                    }
                })
        
        print(f"Document found, fields: {len(document.get('fields', []))}")
        
        # Get the original file path (we'll need to reconstruct this)
        # For now, we'll use the document data directly
        result = {
            'extracted_text': document.get('extracted_text', '')
        }
        
        print(f"Extracted text length: {len(result['extracted_text'])}")
        
        # Trigger automatic training
        training_results = auto_train_from_document(document.get('filename', 'unknown'), document, result)
        
        if training_results:
            print(f"Training successful: {training_results}")
            return Response({
                'success': True,
                'message': 'Training completed successfully',
                'training': training_results
            })
        else:
            print("No training data could be extracted")
            return Response({
                'success': False,
                'message': 'No training data could be extracted'
            })
        
    except Exception as e:
        print(f"Training error: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def get_training_stats(request):
    """Get training statistics"""
    try:
        if not universal_processor:
            return Response({'error': 'Universal Document Processor not available'}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        stats = universal_processor.get_system_stats()
        
        # Add persistent storage info
        persistent_stats = training_storage.get_training_stats()
        stats.update(persistent_stats)
        
        # Add memory storage info for comparison
        stats['documents_in_memory'] = len(documents_storage)
        stats['memory_document_ids'] = list(documents_storage.keys())
        
        return Response(stats)
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def preview_document(request, doc_id):
    """Preview the original document"""
    try:
        document = documents_storage.get(str(doc_id))
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        file_path = document['file_path']
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # For PDFs, serve the file directly
            return Response({
                'success': True,
                'preview_type': 'pdf',
                'preview_url': f'/media/uploads/{os.path.basename(file_path)}'
            })
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp']:
            # For images, serve the file directly
            return Response({
                'success': True,
                'preview_type': 'image',
                'preview_url': f'/media/uploads/{os.path.basename(file_path)}'
            })
        else:
            # For text files, return the content
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return Response({
                    'success': True,
                    'preview_type': 'text',
                    'content': content
                })
            except:
                return Response({
                    'success': True,
                    'preview_type': 'text',
                    'content': document.get('extracted_text', 'No content available')
                })
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def regenerate_document(request, doc_id):
    """Regenerate the original document with filled fields overlaid"""
    try:
        print(f"Regenerating document: {doc_id}")
        
        document = documents_storage.get(str(doc_id))
        if not document:
            print(f"Document not found: {doc_id}")
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        print(f"Document found: {document.get('filename', 'unknown')}")
        
        file_path = document['file_path']
        print(f"File path: {file_path}")
        
        if not os.path.exists(file_path):
            print(f"File does not exist: {file_path}")
            return Response({'error': 'Original file not found'}, status=status.HTTP_404_NOT_FOUND)
        
        file_ext = os.path.splitext(file_path)[1].lower()
        print(f"File extension: {file_ext}")
        
        # Create output filename
        output_filename = f"filled_{os.path.basename(file_path)}"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        print(f"Output path: {output_path}")
        print(f"Fields count: {len(document.get('fields', []))}")
        
        if file_ext == '.pdf':
            # For PDFs, create a new PDF with filled content
            print("Creating filled PDF...")
            success = create_filled_pdf(file_path, output_path, document['fields'])
        elif file_ext in ['.doc', '.docx']:
            # For Word documents, create a new Word document with filled content
            print("Creating filled Word document...")
            success = create_filled_word(file_path, output_path, document['fields'])
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp']:
            # For images, create a new image with text overlaid
            print("Creating filled image...")
            success = create_filled_image(file_path, output_path, document['fields'])
        else:
            # For text files, create a new text file with filled content
            print("Creating filled text file...")
            success = create_filled_text(file_path, output_path, document['fields'])
        
        print(f"Regeneration success: {success}")
        
        if success:
            return Response({
                'success': True,
                'output_file': output_filename,
                'download_url': f'/media/processed/{output_filename}',
                'message': 'Document regenerated with filled fields'
            })
        else:
            return Response({'error': 'Failed to regenerate document'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
    except Exception as e:
        print(f"Regenerate error: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e), 'traceback': traceback.format_exc()}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def create_filled_pdf(input_path, output_path, fields):
    """Create a filled PDF with text overlaid on the original - handles multi-page PDFs"""
    try:
        import fitz  # PyMuPDF
        
        # Open the original PDF
        doc = fitz.open(input_path)
        
        # Create a new PDF document
        new_doc = fitz.open()
        
        # IMPORTANT: Coordinates from field detection are scaled by 2.0 (Matrix(2.0, 2.0))
        # We need to scale them back to match the original PDF dimensions
        SCALE_FACTOR = 2.0
        
        # Group fields by page for easier processing
        fields_by_page = {}
        for field in fields:
            page_num = field.get('page', 0)
            if page_num not in fields_by_page:
                fields_by_page[page_num] = []
            fields_by_page[page_num].append(field)
        
        print(f"Processing {len(doc)} pages with {len(fields)} total fields")
        print(f"Fields by page: {[(p, len(f)) for p, f in fields_by_page.items()]}")
        
        # Process each page
        for page_num in range(len(doc)):
            page = doc[page_num]
            
            # Create a new page with same dimensions
            new_page = new_doc.new_page(width=page.rect.width, height=page.rect.height)
            
            # Copy the original page content
            new_page.show_pdf_page(page.rect, doc, page_num)
            
            # Get fields for this page
            page_fields = fields_by_page.get(page_num, [])
            print(f"Page {page_num}: {len(page_fields)} fields to fill")
            
            # Add filled content to the fields for this page
            for field in page_fields:
                try:
                    user_content = field.get('user_content', '').strip()
                    if not user_content:
                        continue
                    
                    # Check if this is an AcroForm field (already in PDF coordinates)
                    # or a visually detected field (needs scaling from 2x image)
                    field_id = str(field.get('id', ''))
                    is_acroform = field_id.startswith('acroform_')
                    
                    # Get field coordinates safely
                    x_pos = field.get('x_position', field.get('x', 0))
                    y_pos = field.get('y_position', field.get('y', 0))
                    field_width = field.get('width', 100)
                    field_height = field.get('height', 25)
                    
                    if is_acroform:
                        # AcroForm fields are already in PDF coordinates
                        x = float(x_pos)
                        y = float(y_pos)
                        width = float(field_width)
                        height = float(field_height)
                    else:
                        # Visual detection fields need to be scaled back
                        x = float(x_pos) / SCALE_FACTOR
                        y = float(y_pos) / SCALE_FACTOR
                        width = float(field_width) / SCALE_FACTOR
                        height = float(field_height) / SCALE_FACTOR
                    
                    field_type = str(field.get('field_type', 'text')).lower()
                    content = str(user_content)
                    
                    print(f"  Filling field {field_id} ({'acroform' if is_acroform else 'visual'}) at ({x:.1f}, {y:.1f}) with '{content}'")
                    
                    # Draw field content based on type
                    if field_type == 'checkbox':
                        # For checkboxes, draw a checkmark if checked
                        if content.lower() in ['checked', 'true', 'yes']:
                            checkmark_points = [
                                (x + 1, y + height/2),
                                (x + width/3, y + height - 1),
                                (x + width - 1, y + 1)
                            ]
                            new_page.draw_polyline(checkmark_points, color=(0, 0, 0), width=1)
                        
                    elif field_type == 'radio':
                        # For radio buttons, draw a filled circle if selected
                        if content.lower() in ['selected', 'true', 'yes']:
                            center_x = x + width/2
                            center_y = y + height/2
                            radius = min(width, height) / 3
                            new_page.draw_circle((center_x, center_y), radius, color=(0, 0, 0), width=1)
                            # Fill the circle
                            new_page.draw_circle((center_x, center_y), radius/2, color=(0, 0, 0), fill=(0, 0, 0))
                        
                    elif field_type == 'dropdown':
                        # For dropdowns, show the selected option
                        font_size = min(10, height * 0.7)  # Scale font to field height
                        new_page.insert_text(
                            (x + 2, y + height * 0.75),  # Position text baseline
                            content,
                            fontsize=font_size,
                            color=(0, 0, 0),
                            fontname="helv"
                        )
                        
                    else:
                        # For text fields, insert the text normally
                        font_size = min(11, height * 0.7)  # Scale font to field height
                        
                        # Calculate text position - align with baseline
                        text_x = x + 2
                        text_y = y + height * 0.75  # Position baseline at 75% of field height
                        
                        new_page.insert_text(
                            (text_x, text_y),
                            content,
                            fontsize=font_size,
                            color=(0, 0, 0),
                            fontname="helv"
                        )
                        
                except Exception as field_error:
                    print(f"  Error filling field {field.get('id', 'unknown')}: {field_error}")
                    import traceback
                    traceback.print_exc()
                    continue
        
        # Save the new document
        new_doc.save(output_path)
        new_doc.close()
        doc.close()
        
        print(f"PDF saved successfully to {output_path}")
        return True
        
    except Exception as e:
        print(f"Error creating filled PDF: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_filled_word(input_path, output_path, fields):
    """Create a filled Word document with filled content"""
    try:
        from docx import Document
        
        # Open the original Word document
        doc = Document(input_path)
        
        # Create a mapping of field types to their content
        field_content_map = {}
        for field in fields:
            if field['user_content']:
                field_type = field.get('field_type', 'text').lower()
                content = field['user_content']
                context = field['context']
                
                # Format content based on field type
                if field_type == 'checkbox':
                    if content.lower() == 'checked':
                        field_content_map[context] = "☑"  # Checked checkbox
                    else:
                        field_content_map[context] = "☐"  # Unchecked checkbox
                elif field_type == 'radio':
                    if content.lower() == 'selected':
                        field_content_map[context] = "●"  # Selected radio button
                    else:
                        field_content_map[context] = "○"  # Unselected radio button
                elif field_type == 'dropdown':
                    field_content_map[context] = f"[{content}]"  # Show selected option
                else:
                    field_content_map[context] = content
        
        # Process each paragraph to fill in the fields
        for paragraph in doc.paragraphs:
            text = paragraph.text
            text_lower = text.lower().strip()
            
            # Look for field indicators and fill them
            for field_context, content in field_content_map.items():
                if field_context.lower() in text_lower:
                    # If the paragraph ends with ':' or is empty, add the content
                    if text.strip().endswith(':') or text.strip() == '':
                        paragraph.text = text + f" {content}"
                        break
                    # If there's a blank space after the field name, fill it
                    elif ':' in text and len(text.split(':')[1].strip()) == 0:
                        paragraph.text = text + f" {content}"
                        break
        
        # Process tables as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        text_lower = text.lower().strip()
                        
                        # Look for field indicators and fill them
                        for field_type, content in field_content_map.items():
                            if field_type in text_lower:
                                # If the paragraph ends with ':' or is empty, add the content
                                if text.strip().endswith(':') or text.strip() == '':
                                    paragraph.text = text + f" {content}"
                                    break
                                # If there's a blank space after the field name, fill it
                                elif ':' in text and len(text.split(':')[1].strip()) == 0:
                                    paragraph.text = text + f" {content}"
                                    break
        
        # Save the filled document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating filled Word document: {e}")
        return False

def create_filled_image(input_path, output_path, fields):
    """Create a filled image with text overlaid on the original"""
    try:
        import cv2
        import numpy as np
        from PIL import Image, ImageDraw, ImageFont
        
        # Load the original image
        image = cv2.imread(input_path)
        if image is None:
            return False
        
        # Convert to PIL Image for text drawing
        pil_image = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(pil_image)
        
        # Try to load a font that matches the form
        try:
            # Try different font sizes to match the form
            font = ImageFont.truetype("arial.ttf", 14)
        except:
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 14)
            except:
                font = ImageFont.load_default()
        
        # Overlay filled text on the fields
        for field in fields:
            if field['user_content']:
                x = field['x_position']
                y = field['y_position']
                width = field['width']
                height = field['height']
                
                # Don't cover the field with white - just add the text
                # This preserves the original form appearance
                text = field['user_content']
                
                # Calculate text position (left-aligned in the field)
                text_bbox = draw.textbbox((0, 0), text, font=font)
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                
                # Position text in the field (left-aligned, vertically centered)
                text_x = x + 5  # Small margin from left edge
                text_y = y + (height - text_height) // 2
                
                # Draw the text with a slight background for readability
                # Draw a semi-transparent background
                bg_rect = [text_x - 2, text_y - 1, text_x + text_width + 2, text_y + text_height + 1]
                draw.rectangle(bg_rect, fill=(255, 255, 255, 200))  # Semi-transparent white
                
                # Draw the text
                draw.text((text_x, text_y), text, fill='black', font=font)
        
        # Save the result
        pil_image.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating filled image: {e}")
        return False

def create_filled_text(input_path, output_path, fields):
    """Create a filled text file with filled content"""
    try:
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Create a copy of the original content to preserve structure
        filled_content = content
        
        # Replace blank lines or placeholder text with filled content
        lines = filled_content.split('\n')
        
        # Create a mapping of field types to their content
        field_content_map = {}
        for field in fields:
            if field['user_content']:
                field_content_map[field['context']] = field['user_content']
        
        # Process each line to fill in the fields
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Look for field indicators and fill them
            for field_type, content in field_content_map.items():
                if field_type in line_lower:
                    # If the line ends with ':' or is empty, add the content
                    if line.strip().endswith(':') or line.strip() == '':
                        lines[i] = line + f" {content}"
                        break
                    # If there's a blank space after the field name, fill it
                    elif ':' in line and len(line.split(':')[1].strip()) == 0:
                        lines[i] = line + f" {content}"
                        break
        
        # Join the lines back together
        filled_content = '\n'.join(lines)
        
        # Save the filled content
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(filled_content)
        
        return True
        
    except Exception as e:
        print(f"Error creating filled text: {e}")
        return False
