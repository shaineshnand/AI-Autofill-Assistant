from django.shortcuts import render
from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.http import FileResponse, Http404, JsonResponse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
import re
import os
import json
import uuid
from datetime import datetime
import pytesseract
import cv2
import numpy as np
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from ollama_integration import AIDocumentProcessor, OllamaClient
from PIL import Image
import fitz  # PyMuPDF for PDF processing
# Import Sejda automation modules
try:
    from libreoffice_draw_automation import LibreOfficeDrawAutomation
    SEJDA_SIMPLE_AVAILABLE = True
except ImportError:
    SEJDA_SIMPLE_AVAILABLE = False

# Check if pywinauto is available for Sejda automation
try:
    import pywinauto
    PYWINAUTO_AVAILABLE = True
except ImportError:
    PYWINAUTO_AVAILABLE = False
# Use enhanced field detector for better field recognition
try:
    from enhanced_field_detector import EnhancedFieldDetector, convert_form_fields_to_dict
    EnhancedDocumentProcessor = EnhancedFieldDetector
    print("Using Enhanced Field Detector")
except ImportError:
    try:
        from improved_field_detector import ImprovedFieldDetector, convert_form_fields_to_dict
        EnhancedDocumentProcessor = ImprovedFieldDetector
        print("Using Improved Field Detector")
    except ImportError:
        try:
            from enhanced_document_processor import EnhancedDocumentProcessor, convert_form_fields_to_dict
            print("Using Enhanced Document Processor")
        except ImportError:
            from simple_enhanced_processor import SimpleEnhancedProcessor as EnhancedDocumentProcessor, convert_form_fields_to_dict
            print("Using Simple Enhanced Processor")
from intelligent_field_filler import IntelligentFieldFiller

# Import Universal Document Processor for automatic training
try:
    from universal_document_processor import UniversalDocumentProcessor
    universal_processor = UniversalDocumentProcessor()
    print("Universal Document Processor loaded for automatic training")
except ImportError:
    universal_processor = None
    print("Universal Document Processor not available")

# Configure Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Import persistent training storage
from training_storage import training_storage

# In-memory storage (no database needed)
documents_storage = {}
chat_sessions = {}

def get_stored_document(doc_id):
    # 
    # Get document from persistent storage first, then fall back to memory storage.
    # This ensures documents survive server reloads.
    
    
    doc_id = str(doc_id)
    
    # Try persistent storage first (survives reloads)
    try:
        document = training_storage.load_document(doc_id)
        if document:
            # Also cache in memory for faster access
            documents_storage[doc_id] = document
            return document
    except Exception as e:
        print(f"Warning: Could not retrieve from persistent storage: {e}")
    
    # Fall back to memory storage
    return documents_storage.get(doc_id)

def save_document(document):
    # 
    # Save document to both persistent storage and memory.
    # This ensures documents survive server reloads.
    
    
    doc_id = str(document.get('id'))
    
    # Save to persistent storage (survives reloads)
    try:
        training_storage.save_document(doc_id, document)
    except Exception as e:
        print(f"Warning: Could not save to persistent storage: {e}")
    
    # Also save to memory for faster access
    documents_storage[doc_id] = document

def auto_train_from_document(file_path, document, result):
    # 
    # Automatically train the Universal Document Processor from uploaded document
    
    
    if not universal_processor:
        print("Universal processor not available")
        return None
    
    try:
        print(f"Starting auto training for: {file_path}")
        
        # Extract text from document
        text = result.get('extracted_text', '')
        if not text:
            print("No extracted text found")
            return None
        
        print(f"Text length: {len(text)}")
        
        # Classify document type
        doc_type, confidence = universal_processor.classify_document_type(text)
        print(f"Auto-classified document type: {doc_type} (confidence: {confidence:.2f})")
        
        # Create training data from detected fields
        training_samples = []
        fields = document.get('fields', [])
        print(f"Processing {len(fields)} fields")
        print(f"Document keys: {list(document.keys())}")
        print(f"First few fields: {fields[:3] if fields else 'No fields'}")
        
        for i, field in enumerate(fields):
            context = field.get('context', '')
            confidence = field.get('confidence', 0)
            field_type = field.get('field_type', 'text')
            field_id = field.get('id', '')
            
            # Handle both numeric ID and string field_id
            if isinstance(field_id, int):
                field_id = str(field_id)
            
            # For stored documents, the context field contains the technical field ID
            # Extract it for field type mapping
            technical_field_id = context if context else ''
            
            print(f"Field {i}: context='{context}' confidence={confidence} type={field_type} id='{field_id}'")
            
            # For PDF form fields, use the field ID as context if no other context
            if not context or context.strip() == '':
                if field_id:
                    # For PDF form fields, extract meaningful field name from field ID
                    if field_id.startswith('topmostSubform'):
                        # Extract field identifier from technical ID (e.g., f1_01 from topmostSubform[0].Page1[0].Step1a[0].f1_01[0])
                        field_parts = field_id.split('.')
                        if field_parts:
                            last_part = field_parts[-1]
                            if last_part.startswith('f') and '_' in last_part:
                                context = last_part  # Use f1_01, f1_02, etc. as context
                            else:
                                context = f"Field {i+1}"
                        else:
                            context = f"Field {i+1}"
                    else:
                        context = field_id
                else:
                    # Generate context from field ID for form fields
                    context = f"Field {i+1}"
            
            # Always create training samples for detected fields
            # Map field types to universal types using the technical field ID
            universal_field_type = map_field_type_to_universal(field_type, technical_field_id)
            
            # Create training sample with minimum confidence of 0.5
            training_sample = {
                'text': context,
                'field_type': universal_field_type,
                'document_type': doc_type,
                'context': context,
                'confidence': max(0.5, confidence)  # Ensure minimum confidence
            }
            training_samples.append(training_sample)
            print(f"Added training sample: {training_sample}")
        
        if not training_samples:
            print("No valid training samples found")
            return None
        
        print(f"Created {len(training_samples)} training samples")
        
        # Always add samples to training data
        for sample in training_samples:
            universal_processor.training_data.append(sample)
        
        # Also save to persistent storage
        training_storage.add_training_samples(training_samples)
        
        print(f"Added {len(training_samples)} samples to training data. Total samples: {len(universal_processor.training_data)}")
        
        # Train the model (only if we have enough samples)
        if len(universal_processor.training_data) >= 3:
            print("Training model with all samples...")
            training_results = universal_processor.train_model(universal_processor.training_data)
            print(f"Training results: {training_results}")
            
            return {
                'samples_added': len(training_samples),
                'document_type': doc_type,
                'field_type_accuracy': training_results.get('field_type_accuracy', 0),
                'document_type_accuracy': training_results.get('document_type_accuracy', 0),
                'training_samples': training_results.get('training_samples', len(universal_processor.training_data))
            }
        else:
            return {
                'samples_added': len(training_samples),
                'document_type': doc_type,
                'training_samples': len(universal_processor.training_data),
                'note': f'Added to training data. Total samples: {len(universal_processor.training_data)} (need 3+ for retraining)'
            }
    
    except Exception as e:
        print(f"Error in automatic training: {e}")
        import traceback
        traceback.print_exc()
        return None

def map_field_type_to_universal(field_type, field_id=''):
    # 
    # Map detected field types to universal field types
    
    
    # Ensure field_id is a string
    if isinstance(field_id, int):
        field_id = str(field_id)
    elif not isinstance(field_id, str):
        field_id = str(field_id) if field_id else ''
    mapping = {
        'name': 'name',
        'email': 'email', 
        'phone': 'phone',
        'address': 'address',
        'date': 'date_of_birth',
        'ssn': 'ssn',
        'id_number': 'id_number',
        'income': 'income',
        'bank_account': 'bank_account',
        'credit_score': 'credit_score',
        'assets': 'assets',
        'patient_id': 'patient_id',
        'insurance': 'insurance',
        'medications': 'medications',
        'allergies': 'allergies',
        'case_number': 'case_number',
        'court': 'court',
        'attorney': 'attorney',
        'student_id': 'student_id',
        'gpa': 'gpa',
        'major': 'major',
        'company_name': 'company_name',
        'position': 'position',
        'experience_years': 'experience_years',
        'salary': 'expected_salary'
    }
    
    # Try direct mapping first
    if field_type in mapping:
        return mapping[field_type]
    
    # For PDF form fields, try to infer from field ID
    if field_id:
        field_id_lower = field_id.lower()
        
        # Extract the field identifier part (e.g., f1_01 from topmostSubform[0].Page1[0].Step1a[0].f1_01[0])
        field_identifier = field_id
        if field_id.startswith('topmostSubform'):
            field_parts = field_id.split('.')
            if field_parts:
                field_identifier = field_parts[-1]
        
        # W-4 form specific mappings using field identifiers
        if 'name' in field_id_lower or 'f1_01' in field_identifier:
            return 'name'
        elif 'ssn' in field_id_lower or 'f1_02' in field_identifier:
            return 'ssn'
        elif 'address' in field_id_lower or 'f1_03' in field_identifier:
            return 'address'
        elif 'city' in field_id_lower or 'f1_04' in field_identifier:
            return 'address'
        elif 'state' in field_id_lower or 'f1_05' in field_identifier:
            return 'address'
        elif 'zip' in field_id_lower or 'f1_06' in field_identifier:
            return 'address'
        elif 'phone' in field_id_lower:
            return 'phone'
        elif 'email' in field_id_lower:
            return 'email'
        elif 'date' in field_id_lower:
            return 'date_of_birth'
        elif 'income' in field_id_lower or 'wage' in field_id_lower:
            return 'income'
        elif 'dependent' in field_id_lower:
            return 'number'
        elif 'allowance' in field_id_lower:
            return 'number'
    
    return 'text'

class DocumentProcessor:
    def __init__(self):
        self.blank_spaces = []
        self.extracted_text = ""
        self._enhanced_processor = None
        self._intelligent_filler = None
        self.ai_processor = AIDocumentProcessor()
    
    @property
    def enhanced_processor(self):
        if self._enhanced_processor is None:
            self._enhanced_processor = EnhancedDocumentProcessor()
        return self._enhanced_processor
    
    @property
    def intelligent_filler(self):
        if self._intelligent_filler is None:
            self._intelligent_filler = IntelligentFieldFiller()
        return self._intelligent_filler
        
    def pdf_to_images(self, pdf_path):
        # Convert all PDF pages to images for processing
        try:
            # Open PDF
            pdf_document = fitz.open(pdf_path)
            images = []
            
            # Convert each page to image
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                
                # Convert to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to OpenCV format
                nparr = np.frombuffer(img_data, np.uint8)
                image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                images.append((page_num, image))
            
            pdf_document.close()
            return images
        except Exception as e:
            raise Exception(f"Error converting PDF to images: {str(e)}")
    
    def process_document(self, file_path):
        # Process uploaded document using enhanced field detection
        try:
            # Use enhanced processor for better field detection
            result = self.enhanced_processor.process_document(file_path)
            
            # Convert FormField objects to dictionary format
            self.blank_spaces = convert_form_fields_to_dict(result['fields'])
            self.extracted_text = result['extracted_text']
            
            return {
                'extracted_text': self.extracted_text,
                'blank_spaces': self.blank_spaces,
                'total_blanks': len(self.blank_spaces)
            }
        except Exception as e:
            # Fallback to original method if enhanced processor fails
            print(f"Enhanced processor failed, falling back to original method: {e}")
            import traceback
            traceback.print_exc()
            return self._process_document_fallback(file_path)
    
    def _process_document_fallback(self, file_path):
        # Fallback document processing method - handles multi-page PDFs
        try:
            # Check file extension
            file_ext = os.path.splitext(file_path)[1].lower()
            
            all_blank_spaces = []
            all_extracted_text = []
            
            if file_ext == '.pdf':
                # Process all PDF pages
                images = self.pdf_to_images(file_path)
                
                for page_num, image in images:
                    # Convert to grayscale
                    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                    
                    # Apply OCR to extract text
                    page_text = pytesseract.image_to_string(gray)
                    all_extracted_text.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    
                    # Find blank spaces (white rectangles)
                    blank_spaces = self.find_blank_spaces(gray, page_num)
                    
                    # If no blank spaces found, create virtual fields based on text analysis
                    if not blank_spaces:
                        blank_spaces = self.create_virtual_fields_from_text(page_text, gray, page_num)
                    
                    all_blank_spaces.extend(blank_spaces)
                
                self.extracted_text = '\n'.join(all_extracted_text)
                self.blank_spaces = all_blank_spaces
            else:
                # Process image file (single page)
                image = cv2.imread(file_path)
                
                # Check if image was loaded successfully
                if image is None:
                    raise ValueError(f"Could not load image from {file_path}. Please ensure the file is a valid image or PDF format.")
                
                # Convert to grayscale
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                
                # Apply OCR to extract text
                self.extracted_text = pytesseract.image_to_string(gray)
                
                # Find blank spaces (white rectangles)
                self.blank_spaces = self.find_blank_spaces(gray, 0)
                
                # If no blank spaces found, create virtual fields based on text analysis
                if not self.blank_spaces:
                    self.blank_spaces = self.create_virtual_fields_from_text(self.extracted_text, gray, 0)
            
            return {
                'extracted_text': self.extracted_text,
                'blank_spaces': self.blank_spaces,
                'total_blanks': len(self.blank_spaces)
            }
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def find_blank_spaces(self, gray_image, page_num=0):
        # Find blank spaces in the document
        try:
            # Apply adaptive threshold for better edge detection
            thresh = cv2.adaptiveThreshold(gray_image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY_INV, 11, 2)
            
            # Find contours
            contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            blank_spaces = []
            image_height, image_width = gray_image.shape
            
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                area = w * h
                
                # More specific filtering for form fields
                # Look for rectangular areas that could be form fields
                aspect_ratio = w / h if h > 0 else 0
                
                # Filter for reasonable form field sizes and shapes
                if (area > 1000 and area < 100000 and  # Larger minimum size to avoid artifacts
                    w > 50 and h > 20 and  # Larger minimum dimensions
                    w < image_width * 0.8 and h < image_height * 0.3 and  # Not too large
                    (aspect_ratio > 0.5 and aspect_ratio < 10)):  # Reasonable aspect ratio
                    
                    # Check if this area is actually blank (mostly white)
                    roi = gray_image[y:y+h, x:x+w]
                    if roi.size > 0:
                        mean_intensity = np.mean(roi)
                        if mean_intensity > 200:  # Mostly white/blank area
                            blank_spaces.append({
                                'x': int(x),
                                'y': int(y),
                                'width': int(w),
                                'height': int(h),
                                'area': int(area),
                                'context': analyze_context(gray_image, x, y, w, h, self.extracted_text),
                                'page': page_num
                            })
            
            # If no blank spaces found with the above method, try a simpler approach
            if not blank_spaces:
                # Look for white rectangular regions
                _, thresh_white = cv2.threshold(gray_image, 240, 255, cv2.THRESH_BINARY)
                contours_white, _ = cv2.findContours(thresh_white, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
                
                for contour in contours_white:
                    x, y, w, h = cv2.boundingRect(contour)
                    area = w * h
                    
                    if (area > 1000 and area < 50000 and
                        w > 50 and h > 20 and
                        w < image_width * 0.6 and h < image_height * 0.2):
                        
                        blank_spaces.append({
                            'x': int(x),
                            'y': int(y),
                            'width': int(w),
                            'height': int(h),
                            'area': int(area),
                            'context': analyze_context(gray_image, x, y, w, h, self.extracted_text),
                            'page': page_num
                        })
            
            return blank_spaces
            
        except Exception as e:
            print(f"Error in find_blank_spaces: {e}")
            return []
    
    def create_virtual_fields_from_text(self, text, gray_image, page_num=0):
        # Create virtual form fields based on text analysis
        virtual_fields = []
        lines = text.split('\n')
        
        # Enhanced form field patterns with more specific matching
        field_patterns = {
            'name': ['enter your name', 'please enter your name', 'name of dependent', 'full name'],
            'age': ['age of dependent', 'age:', 'years old'],
            'dropdown': ['select an item', 'dropdown', 'combo', 'choose'],
            'checkbox': ['check all that apply', 'option 1', 'option 2', 'option 3'],
            'email': ['email', 'e-mail', 'email address'],
            'phone': ['phone', 'telephone', 'tel', 'mobile'],
            'address': ['address', 'street', 'location'],
            'date': ['date', 'birth', 'dob'],
            'signature': ['signature', 'sign', 'initial']
        }
        
        # Get image dimensions if available
        if gray_image is not None:
            height, width = gray_image.shape
        else:
            height, width = 800, 600  # Default dimensions
        
        # Create virtual fields based on text content
        field_id = 0
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Skip empty lines
            if not line_lower:
                continue
                
            # Check if line contains form field indicators
            field_type = None
            for ftype, patterns in field_patterns.items():
                for pattern in patterns:
                    if pattern in line_lower:
                        field_type = ftype
                        break
                if field_type:
                    break
            
            # If we found a field type, create a virtual field
            if field_type:
                virtual_fields.append({
                    'x': 50,  # Default position
                    'y': 100 + (field_id * 60),  # Spread vertically with more space
                    'width': 400,  # Wider default width
                    'height': 40,   # Taller default height
                    'area': 16000,   # Larger default area
                    'context': field_type,
                    'page': page_num
                })
                field_id += 1
        
        # If still no fields found, create some default ones based on common form elements
        if not virtual_fields:
            # Look for lines that end with colons (common in forms)
            for i, line in enumerate(lines):
                if line.strip().endswith(':') and len(line.strip()) > 3:
                    field_type = 'general'
                    line_lower = line.lower().strip()
                    
                    # Determine field type
                    for ftype, patterns in field_patterns.items():
                        for pattern in patterns:
                            if pattern in line_lower:
                                field_type = ftype
                                break
                    
                    virtual_fields.append({
                        'x': 50,
                        'y': 100 + (len(virtual_fields) * 60),
                        'width': 400,
                        'height': 40,
                        'area': 16000,
                        'context': field_type,
                        'page': page_num
                    })
        
        return virtual_fields
    
    def process_word_document(self, file_path):
        # Process Word documents (.doc, .docx)
        try:
            from docx import Document
            
            # Extract text from Word document
            doc = Document(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            
            # Create virtual fields based on text content
            virtual_fields = self.create_virtual_fields_from_text(content, None)
            
            return {
                'extracted_text': content,
                'blank_spaces': virtual_fields,
                'total_blanks': len(virtual_fields)
            }
        except Exception as e:
            # Fallback to simple text extraction if python-docx fails
            try:
                with open(file_path, 'rb') as f:
                    content = f.read().decode('utf-8', errors='ignore')
                
                virtual_fields = self.create_virtual_fields_from_text(content, None)
                
                return {
                    'extracted_text': content,
                    'blank_spaces': virtual_fields,
                    'total_blanks': len(virtual_fields)
                }
            except Exception as e2:
                raise Exception(f"Error processing Word document: {str(e2)}")
    
    def process_text_document(self, file_path):
        # Process text documents (.txt, .rtf)
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # Create virtual fields based on text content
            virtual_fields = self.create_virtual_fields_from_text(content, None)
            
            return {
                'extracted_text': content,
                'blank_spaces': virtual_fields,
                'total_blanks': len(virtual_fields)
            }
        except Exception as e:
            raise Exception(f"Error processing text document: {str(e)}")

def index(request):
    # Main page view
    document = None
    ollama_status = {'running': False, 'models': []}
    
    # Check if there's a document in session
    if 'current_document_id' in request.session:
        doc_id = request.session['current_document_id']
        document = get_stored_document(doc_id)
        if not document:
            del request.session['current_document_id']
    
    # Check Ollama status
    try:
        ollama = OllamaClient()
        ollama_status = {
            'running': ollama.is_ollama_running(),
            'models': ollama.list_models() if ollama.is_ollama_running() else []
        }
    except:
        pass
    
    context = {
        'document': document,
        'ollama_status': ollama_status,
        'now': datetime.now()
    }
    
    return render(request, 'index.html', context)

def analyze_context(image, x, y, w, h, full_text=""):
    # Analyze context around blank space to suggest content
    # If we have the full extracted text, use it for better analysis
    if full_text:
        lines = full_text.split('\n')
        
        # Look for form field indicators in the text
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Check for specific field types
            if any(keyword in line_lower for keyword in ['enter your name', 'please enter your name', 'name:']):
                return 'name'
            elif any(keyword in line_lower for keyword in ['dependent', 'name of dependent']):
                return 'name'
            elif any(keyword in line_lower for keyword in ['age', 'age of dependent']):
                return 'age'
            elif any(keyword in line_lower for keyword in ['select', 'dropdown', 'combo']):
                return 'dropdown'
            elif any(keyword in line_lower for keyword in ['check', 'option']):
                return 'checkbox'
            elif any(keyword in line_lower for keyword in ['email', 'e-mail']):
                return 'email'
            elif any(keyword in line_lower for keyword in ['phone', 'telephone', 'tel']):
                return 'phone'
            elif any(keyword in line_lower for keyword in ['address']):
                return 'address'
            elif any(keyword in line_lower for keyword in ['date', 'birth']):
                return 'date'
            elif any(keyword in line_lower for keyword in ['signature', 'sign']):
                return 'signature'
    
    # Fallback to OCR on context region
    try:
        padding = 50
        y1 = max(0, y - padding)
        y2 = min(image.shape[0], y + h + padding)
        x1 = max(0, x - padding)
        x2 = min(image.shape[1], x + w + padding)
        
        context_region = image[y1:y2, x1:x2]
        context_text = pytesseract.image_to_string(context_region)
        context_lower = context_text.lower()
        
        if 'name' in context_lower:
            return 'name'
        elif 'address' in context_lower:
            return 'address'
        elif 'phone' in context_lower or 'tel' in context_lower:
            return 'phone'
        elif 'email' in context_lower:
            return 'email'
        elif 'date' in context_lower:
            return 'date'
        elif 'signature' in context_lower:
            return 'signature'
    except:
        pass
    
    return 'general'

@api_view(['POST'])
@csrf_exempt
def upload_fillable_pdf(request):
    # 
    # Upload a fillable PDF that was already processed by Sejda Desktop.
    # This endpoint extracts AcroForm fields and prepares them for AI filling.
    
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method allowed'}, status=405)
    
    if 'file' not in request.FILES:
        return JsonResponse({'error': 'No file uploaded'}, status=400)
    
    uploaded_file = request.FILES['file']
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    if file_ext not in ['.pdf']:
        return JsonResponse({'error': 'Only PDF files are supported'}, status=400)
    
    try:
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Save uploaded fillable PDF
        file_path = f"uploads/{doc_id}_{uploaded_file.name}"
        with default_storage.open(file_path, 'wb') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)
        
        full_file_path = os.path.join(settings.MEDIA_ROOT, file_path)
        
        print(f"Processing fillable PDF from Sejda Desktop: {uploaded_file.name}")
        
        # Open PDF and extract AcroForm fields
        doc = fitz.open(full_file_path)
        fields = []
        field_count = 0
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            widgets = page.widgets()
            
            print(f"Page {page_num}: Found {len(widgets) if widgets else 0} form fields")
            
            if widgets:
                for widget in widgets:
                    field_name = widget.field_name or f"field_{field_count}"
                    field_rect = widget.rect
                    
                    # Create field data
                    field_data = {
                        'id': f"acroform_{field_name}_{page_num}_{field_count}",
                        'field_type': 'text',  # Default type
                        'x_position': int(field_rect.x0),
                        'y_position': int(field_rect.y0),
                        'width': int(field_rect.width),
                        'height': int(field_rect.height),
                        'context': f"Form field: {field_name}",
                        'page': page_num,
                        'page_number': page_num,
                        'user_content': widget.field_value or '',
                        'ai_content': '',
                        'ai_suggestion': '',
                        'ai_enhanced': False,
                        'sejda_field_name': field_name,
                        'detection_method': 'sejda_acroform'
                    }
                    fields.append(field_data)
                    field_count += 1
                    
                    print(f"  Field: {field_name} at ({field_rect.x0:.1f}, {field_rect.y0:.1f})")
        
        doc.close()
        
        print(f"Extracted {len(fields)} AcroForm fields from Sejda Desktop PDF")
        
        # Create document record
        document = {
            'id': doc_id,
            'filename': uploaded_file.name,
            'file_path': file_path,
            'fillable_pdf_path': file_path,  # This PDF is already fillable
            'fields': fields,
            'upload_time': datetime.now().isoformat(),
            'total_fields': len(fields),
            'sejda_processed': True,
            'is_fillable': True
        }
        
        # Save document to storage
        documents_storage.save_document(doc_id, document)
        training_storage.save_document(doc_id, document)
        
        # Store document ID in session
        if hasattr(request, 'session'):
            request.session['current_document_id'] = str(doc_id)
        
        print(f"Document ready for AI filling with {len(fields)} fields")
        
        return JsonResponse({
            'document_id': doc_id,
            'filename': uploaded_file.name,
            'total_fields': len(fields),
            'fields': fields,
            'fillable_pdf_url': f"/media/{file_path}",
            'sejda_processed': True,
            'is_fillable': True,
            'success': True
        })
        
    except Exception as e:
        import traceback
        print(f"Error processing fillable PDF: {e}")
        print(traceback.format_exc())
        return JsonResponse({'error': f'Upload failed: {str(e)}'}, status=500)

@api_view(['POST'])
@csrf_exempt
def upload_document_with_sejda(request):
    # 
    # Upload document and create fillable PDF using Sejda integration
    # This provides better field detection than our custom implementation
    
    
    if request.method != 'POST':
        return JsonResponse({'error': 'Only POST method allowed'}, status=405)
    
    if 'file' not in request.FILES:
        return JsonResponse({'error': 'No file uploaded'}, status=400)
    
    uploaded_file = request.FILES['file']
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    
    if file_ext not in ['.pdf']:
        return JsonResponse({'error': 'Only PDF files are supported'}, status=400)
    
    try:
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Save uploaded file
        file_path = f"uploads/{doc_id}_{uploaded_file.name}"
        with default_storage.open(file_path, 'wb') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)
        
        full_file_path = os.path.join(settings.MEDIA_ROOT, file_path)
        
        # Create Sejda integration instance
        sejda = SejdaIntegration()
        
        # Create fillable PDF using Sejda
        fillable_pdf_path = os.path.join(settings.MEDIA_ROOT, f"processed/fillable_{doc_id}_{uploaded_file.name}")
        os.makedirs(os.path.dirname(fillable_pdf_path), exist_ok=True)
        
        print(f"Creating fillable PDF with Sejda integration...")
        success = sejda.create_fillable_pdf(full_file_path, fillable_pdf_path)
        
        if not success:
            return JsonResponse({'error': 'Failed to create fillable PDF'}, status=500)
        
        # Get form fields from the fillable PDF
        form_fields = sejda.get_form_fields(fillable_pdf_path)
        
        # Convert form fields to our format
        fields = []
        for i, field in enumerate(form_fields):
            field_data = {
                'id': f"sejda_field_{i}",
                'field_type': 'text',  # Default type
                'x_position': field['rect'].x0,
                'y_position': field['rect'].y0,
                'width': field['rect'].width,
                'height': field['rect'].height,
                'context': f"Sejda detected field: {field['name']}",
                'page': field['page'],
                'user_content': '',
                'ai_suggestion': '',
                'ai_enhanced': False,
                'sejda_field_name': field['name']  # Store original Sejda field name
            }
            fields.append(field_data)
        
        # Create document record
        document = {
            'id': doc_id,
            'filename': uploaded_file.name,
            'file_path': file_path,
            'fillable_pdf_path': f"processed/fillable_{doc_id}_{uploaded_file.name}",
            'fields': fields,
            'upload_time': datetime.now().isoformat(),
            'total_fields': len(fields),
            'sejda_processed': True
        }
        
        # Save document to storage
        documents_storage.save_document(doc_id, document)
        training_storage.save_document(doc_id, document)
        
        # Store document ID in session
        if hasattr(request, 'session'):
            request.session['current_document_id'] = str(doc_id)
        
        print(f"Sejda integration created {len(fields)} form fields")
        
        return JsonResponse({
            'document_id': doc_id,
            'filename': uploaded_file.name,
            'total_fields': len(fields),
            'fields': fields,
            'fillable_pdf_url': f"/media/processed/fillable_{doc_id}_{uploaded_file.name}",
            'sejda_processed': True
        })
        
    except Exception as e:
        print(f"Error in Sejda upload: {e}")
        return JsonResponse({'error': f'Upload failed: {str(e)}'}, status=500)

@api_view(['POST'])
def ai_fill_html_document(request, doc_id):
    """Fill HTML document with AI data"""
    try:
        # Get document from storage
        document = get_stored_document(doc_id)
        if not document or not document.get('html_processed'):
            return Response({'error': 'Document not found or not HTML processed'}, status=404)
        
        # Generate AI data for all fields
        from chat.views import IntelligentFieldFiller
        intelligent_filler = IntelligentFieldFiller()
        
        doc_context = {
            'document_type': document.get('document_type', 'form'),
            'total_blanks': document.get('total_blanks', 0),
            'field_types': [field.get('field_type', 'text') for field in document.get('fields', [])],
            'extracted_text': document.get('extracted_text', '')
        }
        
        ai_data = {}
        for field in document.get('fields', []):
            try:
                field_data = {
                    'id': field.get('id', ''),
                    'name': field.get('name', ''),
                    'field_type': field.get('field_type', 'text'),
                    'placeholder': field.get('placeholder', '')
                }
                
                suggested_content = intelligent_filler.generate_field_content(field_data, doc_context)
                ai_data[field.get('id', '')] = suggested_content
                
            except Exception as e:
                print(f"Failed to generate AI data for {field.get('id', '')}: {e}")
                ai_data[field.get('id', '')] = "Sample Data"
        
        # Fill HTML with AI data
        from html_pdf_processor import HTMLPDFProcessor
        processor = HTMLPDFProcessor()
        
        # Read original HTML
        with open(document['html_path'], 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Fill with AI data
        filled_html = processor.fill_html_with_ai_data(html_content, ai_data)
        
        # Save filled HTML
        filled_html_path = document['html_path'].replace('.html', '_filled.html')
        with open(filled_html_path, 'w', encoding='utf-8') as f:
            f.write(filled_html)
        
        # Update document data
        document['ai_filled'] = True
        document['filled_html_path'] = filled_html_path
        document['filled_html_url'] = f"/media/processed/{os.path.basename(filled_html_path)}"
        document['ai_data'] = ai_data
        
        save_document(doc_id, document)
        
        return Response({
            'success': True,
            'message': 'AI data filled successfully',
            'filled_html_url': document['filled_html_url'],
            'ai_data': ai_data
        })
        
    except Exception as e:
        print(f"AI fill error: {e}")
        return Response({'error': str(e)}, status=500)

@api_view(['POST'])
def generate_pdf_from_html(request, doc_id):
    """Generate PDF from filled HTML"""
    try:
        # Get document from storage
        document = get_stored_document(doc_id)
        if not document or not document.get('ai_filled'):
            return Response({'error': 'Document not found or not AI filled'}, status=404)
        
        # Generate PDF from filled HTML
        from html_pdf_processor import HTMLPDFProcessor
        processor = HTMLPDFProcessor()
        
        # Read filled HTML
        with open(document['filled_html_path'], 'r', encoding='utf-8') as f:
            filled_html = f.read()
        
        # Generate output PDF path
        output_filename = f"ai_filled_{document['filename']}"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        
        # Convert HTML to PDF
        processor.html_to_pdf(filled_html, output_path)
        
        # Update document with PDF info
        document['output_pdf_path'] = output_path
        document['output_pdf_url'] = f"/media/processed/{output_filename}"
        
        save_document(doc_id, document)
        
        return Response({
            'success': True,
            'message': 'PDF generated successfully',
            'pdf_url': document['output_pdf_url']
        })
        
    except Exception as e:
        print(f"PDF generation error: {e}")
        return Response({'error': str(e)}, status=500)

@api_view(['POST'])
def upload_document(request):
    # Handle document upload and processing
    try:
        print("=== UPLOAD STARTING ===")
        print("Starting upload_document function")
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Always use HTML-based processing
        use_html_processing = True
        print(f"HTML Processing: Always enabled")
        
        file = request.FILES['file']
        if not file.name:
            return Response({'error': 'No file selected'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Check file size (limit to 50MB)
        if file.size > 50 * 1024 * 1024:
            return Response({'error': 'File too large. Maximum size is 50MB.'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get file extension
        file_ext = os.path.splitext(file.name)[1].lower()
        
        # Only support PDF files for HTML workflow
        if file_ext != '.pdf':
            return Response({'error': 'Only PDF files are supported'}, status=400)
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())
        
        # Save uploaded file
        filename = f"{doc_id}_{file.name}"
        upload_path = os.path.join(settings.MEDIA_ROOT, 'uploads')
        os.makedirs(upload_path, exist_ok=True)
        filepath = os.path.join(upload_path, filename)
        
        with open(filepath, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        
        # HTML-based PDF processing only
        if file_ext == '.pdf':
            print("Using HTML-based PDF processing...")
            try:
                from html_pdf_processor import HTMLPDFProcessor
                processor = HTMLPDFProcessor()
                
                # Step 1: Extract PDF layout and convert to HTML
                layout = processor.extract_pdf_layout(filepath)
                html_content = processor.create_html_template(layout)
                
                # Save HTML content for display
                html_filename = f"html_{doc_id}_{os.path.basename(filepath)}.html"
                html_path = os.path.join(settings.MEDIA_ROOT, 'processed', html_filename)
                os.makedirs(os.path.dirname(html_path), exist_ok=True)
                
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                # Store document data with HTML info
                document_data = {
                    'id': doc_id,
                    'filename': file.name,
                    'file_path': filepath,
                    'total_blanks': len(layout.fields),
                    'fields': layout.fields,
                    'html_path': html_path,
                    'html_url': f"/media/processed/{html_filename}",
                    'document_type': layout.document_type,
                    'extracted_text': layout.extracted_text,
                    'html_processed': True,
                    'ai_filled': False
                }
                
                # Save to storage
                save_document(doc_id, document_data)
                
                print(f"HTML conversion successful: {len(layout.fields)} fields detected")
                return Response({
                    'success': True,
                    'message': 'PDF converted to HTML successfully',
                    'doc_id': doc_id,
                    'html_url': document_data['html_url'],
                    'total_blanks': len(layout.fields),
                    'document_type': layout.document_type,
                    'html_processed': True
                })
                    
            except Exception as e:
                print(f"HTML processing error: {e}")
                return Response({'error': f'HTML processing failed: {str(e)}'}, status=500)
        else:
            return Response({'error': 'Only PDF files are supported'}, status=400)
        
        # Process document based on file type
        processor = DocumentProcessor()
        
        # Process document using enhanced processor (fallback if Sejda not available)
        result = processor.process_document(filepath)
        
        # Create document record in memory
        document = {
            'id': doc_id,
            'filename': file.name,
            'file_path': filepath,
            'extracted_text': result['extracted_text'],
            'total_blanks': result['total_blanks'],
            'uploaded_at': datetime.now().isoformat(),
            'status': 'processed',
            'fields': []
        }
        
        # Store document ID in session (if available)
        if hasattr(request, 'session'):
            request.session['current_document_id'] = str(doc_id)
        
        # For CLEAN Sejda workflow, we don't extract fields during upload
        # Fields will be detected and filled directly in Sejda when user clicks "AI Fill"
        
        # Create field records in memory with intelligent filtering
        # Prioritize: Sejda > AcroForm > Dotted Lines > Visual Blank > Other Visual Detection
        acroform_fields = []
        visual_fields = []
        dots_fields = []
        visual_blank_fields = []
        
        for i, space in enumerate(result['blank_spaces']):
            field = {
                'id': str(space.get('id', i)),
                'field_type': space.get('field_type', space.get('context', 'text')),
                'x_position': space.get('x_position', space.get('x', 0)),
                'y_position': space.get('y_position', space.get('y', 0)),
                'width': space.get('width', 100),
                'height': space.get('height', 25),
                'area': space.get('area', space.get('width', 100) * space.get('height', 25)),
                'context': space.get('context', 'text'),
                'page': space.get('page', 0),  # Include page number
                'user_content': '',
                'ai_suggestion': '',
                'ai_enhanced': False
            }
            
            # Categorize fields by detection method
            field_id = str(space.get('id', i))
            if field_id.startswith('acroform_'):
                acroform_fields.append(field)
            elif field_id.startswith('dots_field_'):
                dots_fields.append(field)
            elif field_id.startswith('visual_blank_'):
                visual_blank_fields.append(field)
            else:
                visual_fields.append(field)
        
        # Add fields in priority order: AcroForm > Dotted Lines > Visual Blank > Other Visual Detection
        # (Sejda CLEAN workflow doesn't need fields during upload - handles them at fill time)
        if acroform_fields:
            document['fields'].extend(acroform_fields)
            document['fields'].extend(dots_fields)  # Always include dotted line fields
            document['fields'].extend(visual_blank_fields)  # Include visual blank fields
            print(f"Added {len(acroform_fields)} AcroForm fields + {len(dots_fields)} dotted line fields + {len(visual_blank_fields)} visual blank fields")
            if visual_fields:
                print(f"Skipped {len(visual_fields)} other visual detection fields (higher priority fields exist)")
        elif dots_fields or visual_blank_fields:
            # No AcroForm fields, but dotted line or visual blank fields exist
            document['fields'].extend(dots_fields)
            document['fields'].extend(visual_blank_fields)
            print(f"Added {len(dots_fields)} dotted line fields + {len(visual_blank_fields)} visual blank fields")
            if visual_fields:
                print(f"Skipped {len(visual_fields)} other visual detection fields (higher priority fields exist)")
        else:
            # No high-priority fields, use all detected fields
            document['fields'].extend(visual_fields)
            print(f"Added {len(visual_fields)} visual detection fields (no high-priority fields found)")
        
        # Automatically create an interactive fillable PDF for PDF uploads
        print(f"Upload processing - File extension: {file_ext}")
        try:
            if file_ext == '.pdf':
                # Create fillable PDF using our method
                print(f"PDF detected - processing fillable PDF creation...")
                processed_dir = os.path.join(settings.MEDIA_ROOT, 'processed')
                os.makedirs(processed_dir, exist_ok=True)
                fillable_filename = f"fillable_{os.path.basename(filepath)}"
                fillable_output_path = os.path.join(processed_dir, fillable_filename)
                print(f"Creating fillable PDF with {len(document['fields'])} fields...")
                success = create_fillable_pdf(filepath, fillable_output_path, document['fields'])
                print(f"Fillable PDF creation result: {success}")
                if success:
                    document['fillable_pdf'] = f"/media/processed/{fillable_filename}"
                    document['fillable_pdf_path'] = fillable_output_path
                    print(f"Fillable PDF created: {fillable_output_path}")
                    
                    # SKIP PyMuPDF filling - Sejda will handle it!
                    # (PyMuPDF fillable PDF is just a backup if Sejda fails)
                    pass  # Skip filling - Sejda does everything!
                else:
                    print(f"Fillable PDF creation failed for: {filepath}")
        except Exception as _auto_fillable_err:
            # Non-fatal: continue without blocking upload
            import traceback
            print(f"Auto fillable PDF generation failed: {_auto_fillable_err}")
            print(f"Traceback: {traceback.format_exc()}")

        # Store in memory
        save_document(document)
        
        # LIBREOFFICE CONVERSION WORKFLOW
        print(f"\nDEBUG: use_libreoffice_conversion = {use_libreoffice_conversion}")
        print(f"DEBUG: pymupdf_auto = {pymupdf_auto}")
        print(f"DEBUG: PYWINAUTO_AVAILABLE = {PYWINAUTO_AVAILABLE}")
        
        if use_libreoffice_conversion and pymupdf_auto:
            print("\n" + "="*60)
            print("STARTING PYMUPDF ACROFORM FILLING!")
            print("="*60)
            
            # Output path for fillable PDF
            fillable_output = filepath.replace('uploads/', 'processed/pymupdf_fillable_')
            fillable_output_full = os.path.join(settings.MEDIA_ROOT, fillable_output)
            
            print(f"Input: {filepath}")
            print(f"Output: {fillable_output_full}")
            
            # FIRST: Generate AI data for all detected fields
            print(f"\nGenerating AI data for {len(document['fields'])} fields...")
            from chat.views import IntelligentFieldFiller
            intelligent_filler = IntelligentFieldFiller()
            
            doc_context = {
                'document_type': 'form',
                'total_blanks': document['total_blanks'],
                'field_types': [field.get('field_type', 'text') for field in document['fields']],
                'extracted_text': document['extracted_text']
            }
            
            # Generate AI data BEFORE opening Sejda
            ai_data_for_libreoffice = {}
            for field in document['fields']:
                try:
                    suggested_content = intelligent_filler.generate_field_content(field, doc_context)
                    ai_data_for_libreoffice[field['id']] = suggested_content
                    print(f"   {field['id']}: '{suggested_content[:30]}...'")
                except Exception as ai_err:
                    print(f"   Failed for {field['id']}: {ai_err}")
            
            print(f"Generated AI data for {len(ai_data_for_libreoffice)} fields")
            print(f"AI Data Details:")
            for field_id, value in ai_data_for_libreoffice.items():
                print(f"   - {field_id}: '{value}'")
            
            # NOW: Fill AcroForm fields using PyMuPDF
            print(f"\nFilling AcroForm fields with PyMuPDF...")
            result_libreoffice = pymupdf_auto.convert_to_fillable(filepath, fillable_output_full, ai_data=ai_data_for_libreoffice, timeout=60)
            
            if result_libreoffice['success']:
                print("\nPyMuPDF AcroForm filling successful!")
                print(f"   Filled {result_libreoffice.get('filled_count', 0)}/{result_libreoffice.get('total_fields', 0)} fields")
                print(f"   Output: {fillable_output_full}")
                
                # Update document with the filled PDF path
                document['pymupdf_fillable_pdf'] = fillable_output
                document['pymupdf_auto_filled'] = True
                
                # Also store AI-filled data in fields
                for field in document['fields']:
                    field_id = field['id']
                    if field_id in ai_data_for_libreoffice:
                        field['ai_content'] = ai_data_for_libreoffice[field_id]
                        field['user_content'] = ai_data_for_libreoffice[field_id]
                
                save_document(document)
                
                print("\n" + "="*60)
                print("AUTOMATIC WORKFLOW COMPLETE!")
                print(f"   Filled PDF ready!")
                print("="*60 + "\n")
                
                # Sejda saves to the input location, not the output location
                sejda_output_path = filepath  # Sejda overwrites the input file
                
                response_data = {
                    'success': True,
                    'document_id': str(doc_id),
                    'document': document,
                    'pymupdf_converted': True,
                    'pymupdf_fillable_url': f"/media/{fillable_output}",  # Use actual PyMuPDF output
                    'auto_download': True,  # FULLY AUTOMATIC - AI fills, saves, opens in new tab
                    'result': {
                        'extracted_text': result['extracted_text'],
                        'total_blanks': len(ai_data_for_libreoffice)
                    },
                    'message': 'PDF filled automatically with LibreOffice Draw! Downloading now...'
                }
                return Response(response_data)
                
                # Skip the field extraction below since we already have everything
                
            else:
                print(f"PyMuPDF conversion failed: {result_libreoffice.get('error', 'Unknown error')}")
                print("   -> Continuing with normal upload")
        
        # Store in memory
        save_document(document)
        
        response_data = {
            'success': True,
            'document_id': str(doc_id),
            'document': document,
            'result': result
        }
        return Response(response_data)
    
    except Exception as e:
        import traceback
        error_trace = traceback.format_exc()
        print(f"Upload error: {e}")
        print(error_trace)
        return Response({'error': str(e), 'traceback': error_trace}, status=500)


        response_data = {
            'success': True,
            'document_id': str(doc_id),
            'document': document,
            'result': {
                'extracted_text': result['extracted_text'],
                'total_blanks': result['total_blanks']
            }
        }
        
        return Response(response_data)
    
    except Exception as e:
        print(f"Upload error: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e), 'traceback': traceback.format_exc()}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def get_document(request, doc_id):
    # Get document information
    try:
        document = get_stored_document(doc_id)
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        return Response(document)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def update_field(request, doc_id):
    # Update a specific field in the document
    try:
        field_id = request.data.get('field_id')
        content = request.data.get('content', '')
        
        document = get_stored_document(doc_id)
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Find and update the field
        for field in document['fields']:
            if field['id'] == field_id:
                field['user_content'] = content
                break
        else:
            return Response({'error': 'Field not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Save to both persistent storage and memory
        save_document(document)
        
        return Response({'success': True})
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def delete_field(request, doc_id):
    # Delete a specific field from the document
    try:
        field_id = request.data.get('field_id')
        
        document = get_stored_document(doc_id)
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Convert field_id to both int and string for flexible matching
        try:
            field_id_int = int(field_id)
            field_id_str = str(field_id)
        except (ValueError, TypeError):
            field_id_int = None
            field_id_str = str(field_id)
        
        # Find and remove the field (match by int OR string)
        original_count = len(document['fields'])
        document['fields'] = [
            field for field in document['fields'] 
            if not (field['id'] == field_id or 
                   field['id'] == field_id_int or 
                   field['id'] == field_id_str or
                   str(field['id']) == field_id_str)
        ]
        
        if len(document['fields']) == original_count:
            return Response({'error': 'Field not found', 'searched_id': field_id, 'total_fields': original_count}, status=status.HTTP_404_NOT_FOUND)
        
        # Update persistent storage
        training_storage.save_document(str(doc_id), document)
        
        return Response({
            'success': True, 
            'remaining_fields': len(document['fields']),
            'deleted_id': field_id
        })
        
    except Exception as e:
        import traceback
        return Response({'error': str(e), 'traceback': traceback.format_exc()}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def generate_pdf(request, doc_id):
    # Generate final PDF with filled content
    try:
        document = documents_storage.get(str(doc_id))
        if not document:
            # Fallback to persistent storage if not found in memory
            print(f"Document {doc_id} not found in memory storage, checking persistent storage...")
            try:
                document = training_storage.load_document(str(doc_id))
                if document:
                    print(f"Document {doc_id} found in persistent storage")
                    # Cache in memory storage
                    documents_storage[str(doc_id)] = document
                else:
                    print(f"Document {doc_id} not found in persistent storage either")
                    return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                print(f"Error retrieving document from persistent storage: {e}")
                return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Prefer the auto-generated fillable PDF as the base
        print(f"Generate PDF - Document keys: {list(document.keys())}")
        print(f"Generate PDF - fillable_pdf_path: {document.get('fillable_pdf_path')}")
        print(f"Generate PDF - file_path: {document.get('file_path')}")
        base_input = document.get('fillable_pdf_path') or document['file_path']
        print(f"Generate PDF - Using base input: {base_input}")
        print(f"Generate PDF - Fillable PDF path available: {document.get('fillable_pdf_path') is not None}")
        processed_path = os.path.join(settings.MEDIA_ROOT, 'processed')
        os.makedirs(processed_path, exist_ok=True)
        output_filename = f"filled_{doc_id}.pdf"
        pdf_path = os.path.join(processed_path, output_filename)

        # Fill either the fillable base (widgets) or draw overlays when needed
        print(f"Generate PDF - Creating filled PDF with {len(document['fields'])} fields...")
        success = create_filled_pdf(base_input, pdf_path, document['fields'])
        if not success and base_input != document['file_path']:
            # Fallback to original file if widget fill failed for any reason
            success = create_filled_pdf(document['file_path'], pdf_path, document['fields'])
        
        if not success:
            return Response({'error': 'Failed to generate filled PDF'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        return Response({
            'success': True,
            'pdf_path': pdf_path,
            'download_url': f"/api/documents/download/{doc_id}"
        })
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def download_pdf(request, doc_id):
    # Download the generated PDF
    try:
        pdf_filename = f"filled_{doc_id}.pdf"
        pdf_path = os.path.join(settings.MEDIA_ROOT, 'processed', pdf_filename)
        
        if not os.path.exists(pdf_path):
            return Response({'error': 'PDF not found'}, status=status.HTTP_404_NOT_FOUND)
        
        return FileResponse(open(pdf_path, 'rb'), as_attachment=True, filename=pdf_filename)
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def clear_session(request):
    # Clear the current document from session
    try:
        # Clear the current document ID from session
        if 'current_document_id' in request.session:
            doc_id = request.session['current_document_id']
            # Optionally remove the document from memory storage
            if str(doc_id) in documents_storage:
                del documents_storage[str(doc_id)]
            del request.session['current_document_id']
        
        return Response({'success': True, 'message': 'Session cleared successfully'})
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def manual_train(request, doc_id):
    # Manually trigger training from a document
    try:
        print(f"Manual training request for document: {doc_id}")
        
        if not universal_processor:
            print("Universal Document Processor not available")
            return Response({'error': 'Universal Document Processor not available'}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Try memory first, then persistent storage
        document = documents_storage.get(str(doc_id))
        
        if not document:
            print(f"Document not found in memory, checking persistent storage: {doc_id}")
            document = training_storage.load_document(str(doc_id))
        
        if not document:
            print(f"Document not found in persistent storage either: {doc_id}")
            # Create training data from existing field patterns
            print("Creating training data from existing field patterns...")
            
            # Add some generic training samples based on common field patterns
            generic_samples = [
                {
                    'text': 'Name',
                    'field_type': 'name',
                    'document_type': 'form',
                    'context': 'Name',
                    'confidence': 0.8
                },
                {
                    'text': 'Email Address',
                    'field_type': 'email',
                    'document_type': 'form',
                    'context': 'Email Address',
                    'confidence': 0.8
                },
                {
                    'text': 'Phone Number',
                    'field_type': 'phone',
                    'document_type': 'form',
                    'context': 'Phone Number',
                    'confidence': 0.8
                },
                {
                    'text': 'Social Security Number',
                    'field_type': 'ssn',
                    'document_type': 'form',
                    'context': 'Social Security Number',
                    'confidence': 0.8
                },
                {
                    'text': 'Address',
                    'field_type': 'address',
                    'document_type': 'form',
                    'context': 'Address',
                    'confidence': 0.8
                }
            ]
            
            # Add samples to training data
            for sample in generic_samples:
                universal_processor.training_data.append(sample)
            
            # Also save to persistent storage
            training_storage.add_training_samples(generic_samples)
            
            print(f"Added {len(generic_samples)} generic training samples")
            
            # Train if we have enough samples
            if len(universal_processor.training_data) >= 3:
                training_results = universal_processor.train_model(universal_processor.training_data)
                return Response({
                    'success': True,
                    'message': 'Training completed with generic samples',
                    'training': {
                        'samples_added': len(generic_samples),
                        'document_type': 'form',
                        'training_samples': len(universal_processor.training_data),
                        'field_type_accuracy': training_results.get('field_type_accuracy', 0),
                        'document_type_accuracy': training_results.get('document_type_accuracy', 0)
                    }
                })
            else:
                return Response({
                    'success': True,
                    'message': 'Added generic training samples',
                    'training': {
                        'samples_added': len(generic_samples),
                        'document_type': 'form',
                        'training_samples': len(universal_processor.training_data),
                        'note': 'Need more samples for model training'
                    }
                })
        
        print(f"Document found, fields: {len(document.get('fields', []))}")
        
        # Get the original file path (we'll need to reconstruct this)
        # For now, we'll use the document data directly
        result = {
            'extracted_text': document.get('extracted_text', '')
        }
        
        print(f"Extracted text length: {len(result['extracted_text'])}")
        
        # Trigger automatic training
        training_results = auto_train_from_document(document.get('filename', 'unknown'), document, result)
        
        if training_results:
            print(f"Training successful: {training_results}")
            return Response({
                'success': True,
                'message': 'Training completed successfully',
                'training': training_results
            })
        else:
            print("No training data could be extracted")
            return Response({
                'success': False,
                'message': 'No training data could be extracted'
            })
        
    except Exception as e:
        print(f"Training error: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def get_training_stats(request):
    # Get training statistics
    try:
        if not universal_processor:
            return Response({'error': 'Universal Document Processor not available'}, 
                          status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        stats = universal_processor.get_system_stats()
        
        # Add persistent storage info
        persistent_stats = training_storage.get_training_stats()
        stats.update(persistent_stats)
        
        # Add memory storage info for comparison
        stats['documents_in_memory'] = len(documents_storage)
        stats['memory_document_ids'] = list(documents_storage.keys())
        
        return Response(stats)
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['GET'])
def preview_document(request, doc_id):
    # Preview the original document
    try:
        document = get_stored_document(doc_id)
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        file_path = document['file_path']
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # For PDFs, serve the file directly
            return Response({
                'success': True,
                'preview_type': 'pdf',
                'preview_url': f'/media/uploads/{os.path.basename(file_path)}'
            })
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp']:
            # For images, serve the file directly
            return Response({
                'success': True,
                'preview_type': 'image',
                'preview_url': f'/media/uploads/{os.path.basename(file_path)}'
            })
        else:
            # For text files, return the content
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                return Response({
                    'success': True,
                    'preview_type': 'text',
                    'content': content
                })
            except:
                return Response({
                    'success': True,
                    'preview_type': 'text',
                    'content': document.get('extracted_text', 'No content available')
                })
        
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def regenerate_document(request, doc_id):
    # Regenerate the original document with filled fields overlaid
    try:
        print(f"Regenerating document: {doc_id}")
        
        document = documents_storage.get(str(doc_id))
        if not document:
            print(f"Document {doc_id} not found in memory storage, checking persistent storage...")
            try:
                document = training_storage.load_document(str(doc_id))
                if document:
                    print(f"Document {doc_id} found in persistent storage")
                    # Cache in memory storage
                    documents_storage[str(doc_id)] = document
                else:
                    print(f"Document {doc_id} not found in persistent storage either")
                    return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
            except Exception as e:
                print(f"Error retrieving document from persistent storage: {e}")
                return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)
        
        print(f"Document found: {document.get('filename', 'unknown')}")
        
        file_path = document['file_path']
        print(f"File path: {file_path}")
        
        if not os.path.exists(file_path):
            print(f"File does not exist: {file_path}")
            return Response({'error': 'Original file not found'}, status=status.HTTP_404_NOT_FOUND)
        
        file_ext = os.path.splitext(file_path)[1].lower()
        print(f"File extension: {file_ext}")
        
        # Prefer the auto-generated fillable PDF as base when available
        print(f"Regenerate - fillable_pdf_path: {document.get('fillable_pdf_path')}")
        print(f"Regenerate - file_path: {file_path}")
        base_input = document.get('fillable_pdf_path') or file_path
        print(f"Regenerate - Using base input: {base_input}")
        
        # Create output filename
        output_filename = f"filled_{os.path.basename(base_input)}"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        print(f"Output path: {output_path}")
        print(f"Fields count: {len(document.get('fields', []))}")
        
        if file_ext == '.pdf':
            # For PDFs, create a new PDF with filled content
            print("Creating filled PDF...")
            success = create_filled_pdf(base_input, output_path, document['fields'])
            if not success and base_input != file_path:
                # Fallback to original file if widget fill fails
                success = create_filled_pdf(file_path, output_path, document['fields'])
        elif file_ext == '.pdf_fillable':
            # Special route (if we ever mark files for interactive conversion)
            print("Creating interactive fillable PDF...")
            success = create_fillable_pdf(file_path, output_path, document['fields'])
        elif file_ext in ['.doc', '.docx']:
            # For Word documents, create a new Word document with filled content
            print("Creating filled Word document...")
            success = create_filled_word(file_path, output_path, document['fields'])
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif', '.webp']:
            # For images, create a new image with text overlaid
            print("Creating filled image...")
            success = create_filled_image(file_path, output_path, document['fields'])
        else:
            # For text files, create a new text file with filled content
            print("Creating filled text file...")
            success = create_filled_text(file_path, output_path, document['fields'])
        
        print(f"Regeneration success: {success}")
        
        if success:
            return Response({
                'success': True,
                'output_file': output_filename,
                'download_url': f'/media/processed/{output_filename}',
                'message': 'Document regenerated with filled fields'
            })
        else:
            return Response({'error': 'Failed to regenerate document'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
    except Exception as e:
        print(f"Regenerate error: {e}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e), 'traceback': traceback.format_exc()}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def create_filled_pdf(input_path, output_path, fields):
    # Create a filled PDF with text overlaid on the original - handles multi-page PDFs
    try:
        import fitz  # PyMuPDF
        
        # Open the original PDF
        doc = fitz.open(input_path)
        
        # We will write directly into the original PDF (keeps AcroForm updates intact)
        
        # IMPORTANT: Enhanced detector renders pages at 3.0x scale for detection
        # Scale coordinates back to match the original PDF dimensions
        SCALE_FACTOR = 3.0
        
        # Group fields by page for easier processing
        fields_by_page = {}
        for field in fields:
            page_num = field.get('page', 0)
            if page_num not in fields_by_page:
                fields_by_page[page_num] = []
            fields_by_page[page_num].append(field)

        # First, fill AcroForm widgets directly inside the original document
        # Build per-page map of acroform field_name -> content
        acro_content_by_page = {}
        print(f"Building AcroForm content map from {len(fields)} total fields")
        for page_num, page_fields in fields_by_page.items():
            print(f"  Page {page_num}: {len(page_fields)} fields")
            for f in page_fields:
                field_id_str = str(f.get('id', ''))
                # Treat both acroform_* and dots_field_* as AcroForm fields
                if field_id_str.startswith('acroform_') or field_id_str.startswith('dots_field_'):
                    # Check both user_content and ai_content
                    user_content = str(f.get('user_content', '')).strip()
                    ai_content = str(f.get('ai_content', '')).strip()
                    content = user_content if user_content else ai_content
                    print(f"    Field {field_id_str}: user_content = '{user_content}', ai_content = '{ai_content}', using = '{content}'")
                    if not content:
                        print(f"    Skipping {field_id_str} (no content)")
                        continue
                    # For dots_field_*, generate the same field name as during widget creation
                    # For acroform_*, use context or field_id_str
                    if field_id_str.startswith('dots_field_'):
                        # Generate the same field name as during widget creation
                        base_name = str(f.get('context', 'text') or 'text').replace(' ', '_')[:32]
                        field_name = f"auto_{base_name}_{page_num}_{field_id_str}"
                    else:
                        field_name = str(f.get('context', '')) or field_id_str.replace('acroform_', '')
                    if page_num not in acro_content_by_page:
                        acro_content_by_page[page_num] = {}
                    acro_content_by_page[page_num][field_name] = content
                    print(f"    Added AcroForm field: {field_name} -> '{content}' on page {page_num}")

        # Apply acroform content
        print(f"Applying AcroForm content to {len(doc)} pages")
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_map = acro_content_by_page.get(page_num, {})
            print(f"Page {page_num}: {len(page_map)} AcroForm fields to fill")
            if page_map:
                print(f"  Fields to fill: {list(page_map.keys())}")
            if not page_map:
                continue
            try:
                widgets = list(page.widgets())
                print(f"Page {page_num}: Found {len(widgets)} widgets")
                for widget in widgets:
                    name = getattr(widget, 'field_name', '') or ''
                    if not name:
                        continue
                    print(f"  Widget: {name}")
                    if name in page_map:
                        content = page_map[name]
                        field_type_str = (getattr(widget, 'field_type_string', '') or '').lower()
                        print(f"    Filling widget {name} with '{content}' (type: {field_type_str})")
                        try:
                            if field_type_str == 'checkbox':
                                # Check if content indicates a checked state
                                checked = str(content).strip().lower() in ['1', 'true', 'yes', 'checked', 'on']
                                widget.set_checked(checked)
                            elif field_type_str in ['radiobutton', 'radio']:
                                # Best effort: mark as selected when content is truthy
                                if str(content).strip():
                                    widget.set_checked(True)
                            else:
                                widget.field_value = content
                            widget.update()
                            print(f"    Successfully filled widget {name}")
                        except Exception as e:
                            print(f"    Failed to fill widget {name}: {e}")
                            pass
                    else:
                        print(f"    Widget {name} not found in page_map")
            except Exception as e:
                print(f"    Error processing page {page_num} widgets: {e}")
                pass
        
        print(f"Processing {len(doc)} pages with {len(fields)} total fields")
        print(f"Fields by page: {[(p, len(f)) for p, f in fields_by_page.items()]}")
        
        # Process each page (draw overlays for non-acro fields directly on original page)
        for page_num in range(len(doc)):
            page = doc[page_num]

            # Get fields for this page
            page_fields = fields_by_page.get(page_num, [])
            print(f"Page {page_num}: {len(page_fields)} fields to fill")
            
            # Add filled content to the fields for this page (only for non-AcroForm fields)
            for field in page_fields:
                try:
                    user_content = field.get('user_content', '').strip()
                    if not user_content:
                        continue
                    
                    # Check if this is an AcroForm field (already in PDF coordinates)
                    # or a visually detected field (needs scaling from 3-times image)
                    field_id = str(field.get('id', ''))
                    is_acroform = field_id.startswith('acroform_')
                    
                    # Skip overlay for AcroForm - they were filled directly above
                    if is_acroform:
                        continue

                    # Get field coordinates safely
                    x_pos = float(field.get('x_position', field.get('x', 0)))
                    y_pos = float(field.get('y_position', field.get('y', 0)))
                    field_width = float(field.get('width', 100))
                    field_height = float(field.get('height', 25))
                    
                    # Visual detection fields need to be scaled back from 3-times detection scale
                    x = float(x_pos) / SCALE_FACTOR
                    y = float(y_pos) / SCALE_FACTOR
                    width = float(field_width) / SCALE_FACTOR
                    height = float(field_height) / SCALE_FACTOR
                    
                    field_type = str(field.get('field_type', 'text')).lower()
                    content = str(user_content)
                    
                    print(f"  Filling field {field_id} ({'acroform' if is_acroform else 'visual'}) at ({x:.1f}, {y:.1f}) with '{content}'")
                    
                    # Draw field content based on type
                    if field_type == 'checkbox':
                        # For checkboxes, draw a checkmark if checked
                        if content.lower() in ['checked', 'true', 'yes']:
                            checkmark_points = [
                                (x + 1, y + height/2),
                                (x + width/3, y + height - 1),
                                (x + width - 1, y + 1)
                            ]
                            page.draw_polyline(checkmark_points, color=(0, 0, 0), width=1)
                        
                    elif field_type == 'radio':
                        # For radio buttons, draw a filled circle if selected
                        if content.lower() in ['selected', 'true', 'yes']:
                            center_x = x + width/2
                            center_y = y + height/2
                            radius = min(width, height) / 3
                            page.draw_circle((center_x, center_y), radius, color=(0, 0, 0), width=1)
                            # Fill the circle
                            page.draw_circle((center_x, center_y), radius/2, color=(0, 0, 0), fill=(0, 0, 0))
                        
                    elif field_type == 'dropdown':
                        # For dropdowns, place text inside the field rectangle using a textbox
                        font_size = max(10, min(14, height * 0.9))
                        text_rect = fitz.Rect(x + 2, y + 2, x + width - 2, y + height - 2)
                        page.insert_textbox(
                            text_rect,
                            content,
                            fontsize=font_size,
                            color=(0, 0, 0),
                            fontname="helv",
                            align=fitz.TEXT_ALIGN_LEFT
                        )
                        
                    else:
                        # For text fields, place content inside the rectangle using a textbox for reliable alignment
                        font_size = max(10, min(14, height * 0.9))
                        text_rect = fitz.Rect(x + 2, y + 2, x + width - 2, y + height - 2)
                        page.insert_textbox(
                            text_rect,
                            content,
                            fontsize=font_size,
                            color=(0, 0, 0),
                            fontname="helv",
                            align=fitz.TEXT_ALIGN_LEFT
                        )
                        
                except Exception as field_error:
                    print(f"  Error filling field {field.get('id', 'unknown')}: {field_error}")
                    import traceback
                    traceback.print_exc()
                    continue
        
        # Save the updated document preserving interactive fields
        doc.save(
            output_path,
            garbage=4,  # Maximum garbage collection
            deflate=True,  # Compress
            clean=True  # Clean up
        )
        doc.close()
        
        print(f"PDF saved successfully to {output_path}")
        return True
        
    except Exception as e:
        print(f"Error creating filled PDF: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_fillable_pdf(input_path, output_path, fields):
    # Create an interactive (AcroForm) fillable PDF based on detected fields.
    # - Adds text and checkbox widgets for non-Acro fields using page coordinates.
    # - Keeps existing AcroForm fields untouched.
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(input_path)

        SCALE_FACTOR = 3.0  # detection was done at 3-times scale

        def _detect_dotted_leaders_pdf_for_widgets(pdf_doc, scale: float = 3.0):
            # Fallback detector: find dotted leaders and yield widget-like field dicts.
            # Coordinates are returned in the same scaled pixel space expected by caller.
            detected = []
            pattern = re.compile(r'(\.{3,}|\.{2,})')
            for pnum in range(len(pdf_doc)):
                try:
                    page = pdf_doc[pnum]
                    raw = page.get_text('rawdict')
                    if not raw or 'blocks' not in raw:
                        continue
                    idx = 0
                    for block in raw.get('blocks', []):
                        for line in block.get('lines', []):
                            spans = line.get('spans', [])
                            if not spans:
                                continue
                            built = ''.join(s.get('text', '') or '' for s in spans)
                            if not built or ('.' not in built and '.' not in built):
                                continue
                            # Build per-char positions
                            char_positions = []
                            for s in spans:
                                t = s.get('text', '') or ''
                                x0, y0, x1, y1 = s.get('bbox', (0, 0, 0, 0))
                                width = max(0.1, x1 - x0)
                                per_char = width / max(1, len(t))
                                cx = x0
                                for ch in t:
                                    char_positions.append((ch, cx, cx + per_char, y0, y1))
                                    cx += per_char
                            if not char_positions:
                                continue
                            for m in pattern.finditer(built):
                                s_idx, e_idx = m.start(), m.end()
                                if s_idx >= len(char_positions):
                                    continue
                                start_char = char_positions[s_idx]
                                end_char = char_positions[min(e_idx - 1, len(char_positions) - 1)]
                                x0 = start_char[1]
                                x1 = end_char[2]
                                y0 = min(start_char[3], end_char[3])
                                y1 = max(start_char[4], end_char[4])

                                est_w = max(80.0, min((x1 - x0) * 2.0, 400.0))
                                est_h = max(20.0, min((y1 - y0) * 1.2, 28.0))
                                rect_x = x0
                                rect_y = y0 - est_h

                                page_w, page_h = float(page.rect.width), float(page.rect.height)
                                sx = max(0.0, min(rect_x, page_w)) * scale
                                sy = max(0.0, min(rect_y, page_h)) * scale
                                sw = max(40.0, min(est_w, page_w - rect_x)) * scale
                                sh = max(16.0, min(est_h, page_h * 0.1)) * scale

                                context_before = built[max(0, s_idx-40):s_idx]
                                context_after = built[e_idx:min(len(built), e_idx+40)]
                                combined = (context_before + ' ' + context_after).lower()
                                if 'day' in combined and 'month' not in combined:
                                    ftype = 'day'
                                elif 'month' in combined:
                                    ftype = 'month'
                                elif 'year' in combined or '20' in combined:
                                    ftype = 'year'
                                elif any(k in combined for k in ['employer', 'employee', 'company']):
                                    ftype = 'name'
                                else:
                                    ftype = 'text'

                                detected.append({
                                    'id': f'dots_pdf_p{pnum}_{idx}',
                                    'field_type': ftype,
                                    'x_position': int(sx),
                                    'y_position': int(sy),
                                    'width': int(sw),
                                    'height': int(sh),
                                    'context': (context_before + ' ' + context_after).strip()[:120] or 'dotted_line',
                                    'page': pnum
                                })
                                idx += 1
                except Exception:
                    continue
            return detected

        def _detect_anchor_fields(pdf_doc, scale: float = 3.0):
            # Heuristic anchors for common contract lines where dots might be vector graphics.
            # - Creates fields near words: 'day', 'month', 'year', 'Employer', 'Employee'
            # Coordinates returned in 3-times image space (scale applied).
            results = []
            keywords = ['day', 'month', 'year', 'employer', 'employee']
            for pnum in range(len(pdf_doc)):
                try:
                    page = pdf_doc[pnum]
                    words = page.get_text('words') or []
                    if not words:
                        continue
                    # Normalize list of tuples: (x0,y0,x1,y1,word,block,line,wordno)
                    for w in words:
                        if len(w) < 5:
                            continue
                        x0, y0, x1, y1, token = float(w[0]), float(w[1]), float(w[2]), float(w[3]), str(w[4]).lower()
                        if token not in keywords:
                            continue
                        # default sizes
                        h = max(16.0, min(24.0, (y1 - y0) * 1.1))
                        if token == 'day' or token == 'month':
                            # field to the left of the word
                            width = 100.0 if token == 'day' else 140.0
                            rx = max(0.0, x0 - (width + 8.0))
                            ry = max(0.0, y0 - (h * 0.2))
                        elif token == 'year':
                            # field to the right of the word
                            width = 120.0
                            rx = x1 + 8.0
                            ry = max(0.0, y0 - (h * 0.2))
                        elif token in ('employer', 'employee'):
                            # long field to the left edge
                            width = 360.0
                            left_margin = max(0.0, page.rect.x0 + 50.0)
                            rx = max(0.0, min(x0 - width - 12.0, x0 - 60.0, left_margin))
                            ry = max(0.0, y0 - (h * 0.2))
                        else:
                            continue

                        # Clamp to page
                        rx = max(0.0, min(rx, float(page.rect.width) - 10.0))
                        ry = max(0.0, min(ry, float(page.rect.height) - h))

                        results.append({
                            'id': f'anchor_{token}_p{pnum}_{int(rx)}',
                            'field_type': 'text' if token not in ('day','month','year') else token,
                            'x_position': int(rx * scale),
                            'y_position': int(ry * scale),
                            'width': int(width * scale),
                            'height': int(h * scale),
                            'context': token,
                            'page': pnum
                        })
                except Exception:
                    continue
            return results

        # Always detect dotted leaders and merge with provided fields
        def _overlaps(a, b) -> bool:
            try:
                ax1 = float(a.get('x_position', a.get('x', 0)))
                ay1 = float(a.get('y_position', a.get('y', 0)))
                ax2 = ax1 + float(a.get('width', 0))
                ay2 = ay1 + float(a.get('height', 0))
                bx1 = float(b.get('x_position', b.get('x', 0)))
                by1 = float(b.get('y_position', b.get('y', 0)))
                bx2 = bx1 + float(b.get('width', 0))
                by2 = by1 + float(b.get('height', 0))
                ix1 = max(ax1, bx1)
                iy1 = max(ay1, by1)
                ix2 = min(ax2, bx2)
                iy2 = min(ay2, by2)
                if ix1 >= ix2 or iy1 >= iy2:
                    return False
                inter = (ix2 - ix1) * (iy2 - iy1)
                a_area = (ax2 - ax1) * (ay2 - ay1)
                b_area = (bx2 - bx1) * (by2 - by1)
                # consider overlapping if > 25% of smaller area
                return inter > 0.25 * min(a_area, b_area)
            except Exception:
                return False

        try:
            existing = list(fields or [])
        except Exception:
            existing = []

        # Use existing dots_field_* and visual_blank_* fields from enhanced detection
        dotted_fields = [f for f in existing if f.get('id', '').startswith('dots_field_')]
        visual_blank_fields = [f for f in existing if f.get('id', '').startswith('visual_blank_')]
        print(f"Found {len(dotted_fields)} existing dotted fields and {len(visual_blank_fields)} visual blank fields to convert to AcroForm widgets")
        
        # Convert dots_field_* and visual_blank_* fields to widget format
        dotted_extra = []
        for field in dotted_fields + visual_blank_fields:
            try:
                widget_field = {
                    'id': field['id'],
                    'field_type': field.get('field_type', 'text'),
                    'x_position': field['x_position'],
                    'y_position': field['y_position'], 
                    'width': field['width'],
                    'height': field['height'],
                    'context': field.get('context', ''),
                    'page': field.get('page_number', 0)
                }
                dotted_extra.append(widget_field)
                print(f"  Converting {field['id']} to widget format")
            except Exception as e:
                print(f"  Error converting {field.get('id', 'unknown')}: {e}")
                continue

        # If no dotted fields from detection, try fallback detection
        if not dotted_extra:
            try:
                dotted_extra = _detect_dotted_leaders_pdf_for_widgets(doc, SCALE_FACTOR)
                print(f"Fallback detection found {len(dotted_extra)} dotted fields")
            except Exception as e:
                print(f"Fallback dotted detection failed: {e}")
                dotted_extra = []

        # If still few or no dotted fields found, add anchor-based fields as safety net
        try:
            if len([f for f in dotted_extra if f.get('page', 0) == 0]) < 2:
                anchor_fields = _detect_anchor_fields(doc, SCALE_FACTOR)
                dotted_extra.extend(anchor_fields)
                print(f"Added {len(anchor_fields)} anchor-based fields as safety net")
        except Exception as e:
            print(f"Anchor field detection failed: {e}")
            pass

        # Merge without duplicates
        for d in dotted_extra:
            dup = False
            for e in existing:
                # same page and overlapping
                if int(d.get('page', 0)) == int(e.get('page', 0)) and _overlaps(d, e):
                    dup = True
                    break
            if not dup:
                existing.append(d)

        fields = existing

        # Group fields by page
        fields_by_page = {}
        for field in fields:
            page_num = field.get('page', 0)
            fields_by_page.setdefault(page_num, []).append(field)

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_fields = fields_by_page.get(page_num, [])
            print(f"Page {page_num}: Creating AcroForm widgets for {len(page_fields)} fields")
            print(f"  Field types: {[str(f.get('id', '')) for f in page_fields]}")

            for f in page_fields:
                try:
                    field_id = str(f.get('id', ''))
                    # Skip if it's already an AcroForm field
                    if field_id.startswith('acroform_'):
                        continue
                    
                    # ONLY create widgets for dotted line fields, not regular text boxes
                    if not field_id.startswith('dots_field_'):
                        print(f"  Skipping {field_id} - not a dotted line field")
                        continue

                    field_type = str(f.get('field_type', 'text')).lower()
                    # For dotted fields, use the x1,y1,x2,y2 coordinates directly
                    # These are now exact PDF coordinates (no scaling needed)
                    if 'x1' in f and 'y1' in f and 'x2' in f and 'y2' in f:
                        x1 = float(f['x1'])
                        y1 = float(f['y1']) 
                        x2 = float(f['x2'])
                        y2 = float(f['y2'])
                        
                        # Use exact PDF coordinates directly (no scaling)
                        x = x1
                        y = y1
                        w = x2 - x1
                        h = y2 - y1
                        print(f"    Using exact PDF coordinates: ({x1:.1f}, {y1:.1f}) to ({x2:.1f}, {y2:.1f})")
                    else:
                        # Fallback to old method if x1,y1,x2,y2 not available
                        x_raw = float(f.get('x_position', f.get('x', 0)))
                        y_raw = float(f.get('y_position', f.get('y', 0)))
                        w_raw = float(f.get('width', 100))
                        h_raw = float(f.get('height', 25))
                        
                        x = x_raw / SCALE_FACTOR
                        y = y_raw / SCALE_FACTOR
                        w = w_raw / SCALE_FACTOR
                        h = h_raw / SCALE_FACTOR
                        print(f"    Using fallback scaled coordinates")
                    rect = fitz.Rect(x, y, x + w, y + h)

                    # Generate field name to keep them unique and readable
                    # Use the field's actual page number, not the loop page number
                    field_page_num = f.get('page', page_num)
                    base_name = str(f.get('context', field_type) or field_type).replace(' ', '_')[:32]
                    field_name = f"auto_{base_name}_{field_page_num}_{field_id}"

                    # Choose widget type (fallback to text when API does not support others)
                    widget_type = fitz.PDF_WIDGET_TYPE_TEXT
                    if field_type == 'checkbox':
                        widget_type = fitz.PDF_WIDGET_TYPE_CHECKBOX

                    # Check API availability
                    print(f"  Creating widget for {field_id} at ({x:.1f}, {y:.1f}) size ({w:.1f}x{h:.1f})")
                    print(f"    Field name: {field_name}")
                    print(f"    Field page: {field_page_num}, Loop page: {page_num}")
                    print(f"    Widget type: {widget_type}")
                    print(f"    Final coordinates: x={x:.1f}, y={y:.1f}, w={w:.1f}, h={h:.1f}")
                    print(f"    Rect: {rect}")
                    print(f"    Context: '{f.get('context', '')}'")
                    print(f"    User content: '{f.get('user_content', '')}'")
                    print(f"    AI content: '{f.get('ai_content', '')}'")
                    print(f"    Field detection method: {f.get('detection_method', 'unknown')}")
                    if hasattr(page, 'add_widget'):
                        print(f"    PyMuPDF add_widget method available")
                        try:
                            # Create a new Widget object
                            widget = fitz.Widget()
                            
                            # Set widget properties
                            widget.rect = rect
                            widget.field_name = field_name
                            widget.field_label = field_name
                            widget.field_type = widget_type
                            # Set initial content (prefer user_content, fallback to ai_content)
                            initial_content = f.get('user_content', '') or f.get('ai_content', '')
                            widget.field_value = initial_content
                            widget.field_flags = 0  # Standard editable field
                            widget.border_color = (0, 0, 1)  # Blue border for visibility
                            widget.border_width = 1
                            widget.fill_color = (0.95, 0.95, 1)  # Light blue background
                            widget.text_color = (0, 0, 0)  # Black text
                            widget.text_font = "helv"  # Helvetica font
                            widget.text_fontsize = 10
                            
                            if widget_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                                # Default unchecked; leave content to user
                                widget.field_value = False
                            
                            # Add the widget to the page
                            page.add_widget(widget)
                            widget.update()
                            print(f"    Widget created and updated successfully")
                        except Exception as widget_error:
                            print(f"    ERROR creating widget: {widget_error}")
                            import traceback
                            traceback.print_exc()
                            # Continue with next widget
                    else:
                        # Fallback: draw a thin rectangle to indicate input area (non-interactive)
                        print(f"    PyMuPDF add_widget method NOT available - using fallback rectangle")
                        page.draw_rect(rect, color=(1, 0, 0), width=2)  # Red rectangle for visibility
                except Exception:
                    # Continue even if a single widget fails
                    continue

        # Save to output with proper PDF settings to preserve forms
        doc.save(
            output_path,
            garbage=4,  # Maximum garbage collection
            deflate=True,  # Compress
            clean=True,  # Clean up
            pretty=False  # Don't pretty-print (smaller file)
        )
        doc.close()
        return True
    except Exception as e:
        print(f"Error creating fillable PDF: {e}")
        import traceback
        traceback.print_exc()
        return False

def create_filled_word(input_path, output_path, fields):
    # Create a filled Word document with filled content
    try:
        from docx import Document
        
        # Open the original Word document
        doc = Document(input_path)
        
        # Create a mapping of field types to their content
        field_content_map = {}
        for field in fields:
            if field['user_content']:
                field_type = field.get('field_type', 'text').lower()
                content = field['user_content']
                context = field['context']
                
                # Format content based on field type
                if field_type == 'checkbox':
                    if content.lower() == 'checked':
                        field_content_map[context] = "[X]"  # Checked checkbox
                    else:
                        field_content_map[context] = "[ ]"  # Unchecked checkbox
                elif field_type == 'radio':
                    if content.lower() == 'selected':
                        field_content_map[context] = "(*)"  # Selected radio button
                    else:
                        field_content_map[context] = "O"  # Unselected radio button
                elif field_type == 'dropdown':
                    field_content_map[context] = f"[{content}]"  # Show selected option
                else:
                    field_content_map[context] = content
        
        # Process each paragraph to fill in the fields
        for paragraph in doc.paragraphs:
            text = paragraph.text
            text_lower = text.lower().strip()
            
            # Look for field indicators and fill them
            for field_context, content in field_content_map.items():
                if field_context.lower() in text_lower:
                    # If the paragraph ends with ':' or is empty, add the content
                    if text.strip().endswith(':') or text.strip() == '':
                        paragraph.text = text + f" {content}"
                        break
                    # If there's a blank space after the field name, fill it
                    elif ':' in text and len(text.split(':')[1].strip()) == 0:
                        paragraph.text = text + f" {content}"
                        break
        
        # Process tables as well
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        text_lower = text.lower().strip()
                        
                        # Look for field indicators and fill them
                        for field_type, content in field_content_map.items():
                            if field_type in text_lower:
                                # If the paragraph ends with ':' or is empty, add the content
                                if text.strip().endswith(':') or text.strip() == '':
                                    paragraph.text = text + f" {content}"
                                    break
                                # If there's a blank space after the field name, fill it
                                elif ':' in text and len(text.split(':')[1].strip()) == 0:
                                    paragraph.text = text + f" {content}"
                                    break
        
        # Save the filled document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating filled Word document: {e}")
        return False

def create_filled_image(input_path, output_path, fields):
    # Create a filled image with text overlaid on the original
    try:
        import cv2
        import numpy as np
        from PIL import Image, ImageDraw, ImageFont
        
        # Load the original image
        image = cv2.imread(input_path)
        if image is None:
            return False
        
        # Convert to PIL Image for text drawing
        pil_image = Image.fromarray(cv2.cvtColor(image, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(pil_image)
        
        # Try to load a font that matches the form
        try:
            # Try different font sizes to match the form
            font = ImageFont.truetype("arial.ttf", 14)
        except:
            try:
                font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 14)
            except:
                font = ImageFont.load_default()
        
        # Overlay filled text on the fields
        for field in fields:
            if field['user_content']:
                x = field['x_position']
                y = field['y_position']
                width = field['width']
                height = field['height']
                
                # Don't cover the field with white - just add the text
                # This preserves the original form appearance
                text = field['user_content']
                
                # Calculate text position (left-aligned in the field)
                text_bbox = draw.textbbox((0, 0), text, font=font)
                text_width = text_bbox[2] - text_bbox[0]
                text_height = text_bbox[3] - text_bbox[1]
                
                # Position text in the field (left-aligned, vertically centered)
                text_x = x + 5  # Small margin from left edge
                text_y = y + (height - text_height) // 2
                
                # Draw the text with a slight background for readability
                # Draw a semi-transparent background
                bg_rect = [text_x - 2, text_y - 1, text_x + text_width + 2, text_y + text_height + 1]
                draw.rectangle(bg_rect, fill=(255, 255, 255, 200))  # Semi-transparent white
                
                # Draw the text
                draw.text((text_x, text_y), text, fill='black', font=font)
        
        # Save the result
        pil_image.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error creating filled image: {e}")
        return False

def create_filled_text(input_path, output_path, fields):
    # Create a filled text file with filled content
    try:
        with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Create a copy of the original content to preserve structure
        filled_content = content
        
        # Replace blank lines or placeholder text with filled content
        lines = filled_content.split('\n')
        
        # Create a mapping of field types to their content
        field_content_map = {}
        for field in fields:
            if field['user_content']:
                field_content_map[field['context']] = field['user_content']
        
        # Process each line to fill in the fields
        for i, line in enumerate(lines):
            line_lower = line.lower().strip()
            
            # Look for field indicators and fill them
            for field_type, content in field_content_map.items():
                if field_type in line_lower:
                    # If the line ends with ':' or is empty, add the content
                    if line.strip().endswith(':') or line.strip() == '':
                        lines[i] = line + f" {content}"
                        break
                    # If there's a blank space after the field name, fill it
                    elif ':' in line and len(line.split(':')[1].strip()) == 0:
                        lines[i] = line + f" {content}"
                        break
        
        # Join the lines back together
        filled_content = '\n'.join(lines)
        
        # Save the filled content
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(filled_content)
        
        return True
        
    except Exception as e:
        print(f"Error creating filled text: {e}")
        return False

@api_view(['POST'])
def make_fillable_pdf(request, doc_id):
    # Convert the uploaded PDF into an Interactive fillable PDF by adding widgets for detected fields.
    try:
        document = get_stored_document(doc_id)
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)

        input_path = document['file_path']
        if not os.path.exists(input_path):
            return Response({'error': 'Original file not found'}, status=status.HTTP_404_NOT_FOUND)

        output_filename = f"fillable_{os.path.basename(input_path)}"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        success = create_fillable_pdf(input_path, output_path, document['fields'])
        if success:
            return Response({
                'success': True,
                'output_file': output_filename,
                'download_url': f'/media/processed/{output_filename}',
                'message': 'Interactive fillable PDF created'
            })
        else:
            return Response({'error': 'Failed to create fillable PDF'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def recreate_editable_pdf(request, doc_id):
    """
    Extract text from PDF and create a new editable PDF with AI auto-filled data.
    This recreates the PDF exactly like the original but makes it editable with AI-filled content.
    """
    try:
        # Import PDF recreator
        from pdf_recreator import PDFRecreator
        
        document = get_stored_document(doc_id)
        if not document:
            return Response({'error': 'Document not found'}, status=status.HTTP_404_NOT_FOUND)

        input_path = document['file_path']
        if not os.path.exists(input_path):
            return Response({'error': 'Original file not found'}, status=status.HTTP_404_NOT_FOUND)

        # Create output filename
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        output_filename = f"{base_name}_recreated_ai_filled.pdf"
        output_path = os.path.join(settings.MEDIA_ROOT, 'processed', output_filename)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        # Initialize PDF recreator and process
        recreator = PDFRecreator()
        result = recreator.process_pdf(input_path, output_path)
        
        if result['success']:
            # Update document record with new output
            document['recreated_pdf_path'] = output_path
            save_document(document)
            
            return Response({
                'success': True,
                'output_file': output_filename,
                'download_url': f'/media/processed/{output_filename}',
                'message': 'PDF recreated with AI-filled data',
                'document_type': result['document_type'],
                'fields_detected': result['fields_detected'],
                'fields_filled': result['fields_filled'],
                'extracted_text_preview': result['extracted_text'],
                'filled_data': result['filled_data']
            })
        else:
            return Response({
                'error': 'Failed to recreate PDF',
                'details': result.get('error', 'Unknown error')
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
    except ImportError as e:
        return Response({
            'error': 'PDF recreator module not available',
            'details': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Exception as e:
        return Response({
            'error': 'Failed to recreate PDF',
            'details': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

