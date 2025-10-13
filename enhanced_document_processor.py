#!/usr/bin/env python
"""
Enhanced Document Processor for AI Autofill Assistant
Provides accurate field detection and positioning for various document types
"""
import os
import cv2
import numpy as np
import pytesseract
import fitz  # PyMuPDF
import re
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from PIL import Image, ImageDraw, ImageFont
import json
import random

@dataclass
class FormField:
    """Represents a form field with precise positioning and metadata"""
    id: str
    field_type: str
    x: int
    y: int
    width: int
    height: int
    context: str
    confidence: float
    is_required: bool = False
    placeholder: str = ""
    validation_pattern: str = ""

class EnhancedDocumentProcessor:
    """Enhanced document processor with improved field detection"""
    
    def __init__(self):
        self.extracted_text = ""
        self.document_type = ""
        self.fields = []
        
    def process_document(self, file_path: str) -> Dict:
        """Process document and return enhanced field detection results"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                return self._process_pdf(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._process_image(file_path)
            elif file_ext in ['.doc', '.docx']:
                return self._process_word(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
                
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def _process_pdf(self, file_path: str) -> Dict:
        """Process PDF with both AcroForm detection and OCR analysis"""
        try:
            # First, try to detect AcroForm fields
            acroform_fields = self._detect_pdf_form_fields(file_path)
            
            # Also perform OCR analysis for additional context
            ocr_fields = self._detect_fields_from_ocr(file_path)
            
            # Combine and deduplicate fields
            combined_fields = self._merge_field_detections(acroform_fields, ocr_fields)
            
            # Extract text for context
            self.extracted_text = self._extract_pdf_text(file_path)
            
            return {
                'extracted_text': self.extracted_text,
                'fields': combined_fields,
                'total_fields': len(combined_fields),
                'document_type': 'pdf',
                'has_acroform': len(acroform_fields) > 0
            }
            
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    def _detect_pdf_form_fields(self, file_path: str) -> List[FormField]:
        """Detect AcroForm fields in PDF"""
        fields = []
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                form_fields = page.widgets()
                
                for i, field in enumerate(form_fields):
                    rect = field.rect
                    field_type = self._classify_pdf_field_type(field)
                    
                    form_field = FormField(
                        id=f"pdf_field_{page_num}_{i}",
                        field_type=field_type,
                        x=int(rect.x0),
                        y=int(rect.y0),
                        width=int(rect.width),
                        height=int(rect.height),
                        context=self._analyze_field_context(field, page),
                        confidence=0.9,  # High confidence for AcroForm fields
                        is_required=field.field_flags & 2 != 0,  # Required flag
                        placeholder=field.field_value or ""
                    )
                    fields.append(form_field)
            
            doc.close()
            return fields
            
        except Exception as e:
            print(f"Error detecting PDF form fields: {e}")
            return []
    
    def _classify_pdf_field_type(self, field) -> str:
        """Classify PDF field type based on field properties"""
        field_type = field.field_type_string.lower()
        
        if field_type == 'text':
            # Analyze field name and value for more specific classification
            field_name = (field.field_name or "").lower()
            field_value = (field.field_value or "").lower()
            
            if any(keyword in field_name for keyword in ['email', 'e-mail']):
                return 'email'
            elif any(keyword in field_name for keyword in ['phone', 'tel', 'mobile']):
                return 'phone'
            elif any(keyword in field_name for keyword in ['name', 'first', 'last']):
                return 'name'
            elif any(keyword in field_name for keyword in ['address', 'street']):
                return 'address'
            elif any(keyword in field_name for keyword in ['date', 'birth', 'dob']):
                return 'date'
            elif any(keyword in field_name for keyword in ['age', 'years']):
                return 'age'
            else:
                return 'text'
        elif field_type == 'checkbox':
            return 'checkbox'
        elif field_type == 'combobox' or field_type == 'listbox':
            return 'dropdown'
        elif field_type == 'signature':
            return 'signature'
        else:
            return field_type
    
    def _analyze_field_context(self, field, page) -> str:
        """Analyze context around PDF field"""
        try:
            # Get text around the field
            rect = field.rect
            expanded_rect = fitz.Rect(
                max(0, rect.x0 - 50),
                max(0, rect.y0 - 20),
                min(page.rect.width, rect.x1 + 50),
                min(page.rect.height, rect.y1 + 20)
            )
            
            context_text = page.get_textbox(expanded_rect).lower()
            
            # Look for field labels and context
            if any(keyword in context_text for keyword in ['name', 'full name', 'enter name']):
                return 'name'
            elif any(keyword in context_text for keyword in ['email', 'e-mail']):
                return 'email'
            elif any(keyword in context_text for keyword in ['phone', 'telephone', 'tel']):
                return 'phone'
            elif any(keyword in context_text for keyword in ['address', 'street']):
                return 'address'
            elif any(keyword in context_text for keyword in ['date', 'birth', 'dob']):
                return 'date'
            elif any(keyword in context_text for keyword in ['age', 'years old']):
                return 'age'
            elif any(keyword in context_text for keyword in ['signature', 'sign']):
                return 'signature'
            else:
                return 'text'
                
        except Exception as e:
            print(f"Error analyzing field context: {e}")
            return 'text'
    
    def _detect_fields_from_ocr(self, file_path: str) -> List[FormField]:
        """Detect form fields using OCR and image analysis - handles all pages"""
        try:
            # Convert PDF to images
            doc = fitz.open(file_path)
            all_fields = []
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                mat = fitz.Matrix(3.0, 3.0)  # Match 3x scaling used elsewhere
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                
                # Convert to OpenCV format
                nparr = np.frombuffer(img_data, np.uint8)
                image = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
                
                # Extract text with bounding boxes
                ocr_data = pytesseract.image_to_data(gray, output_type=pytesseract.Output.DICT)
                
                # Find form fields using advanced image processing
                fields = self._detect_form_fields_advanced(gray, ocr_data, page_num)
                all_fields.extend(fields)

                # Detect dotted-leader placeholders from the page text
                try:
                    # PDF-native: use rawdict to compute accurate positions
                    dotted_fields = self._detect_dotted_leader_fields_pdf(page, page_num, scale=3.0)
                    # Fallback to text heuristic if none found
                    if not dotted_fields:
                        page_text = page.get_text() or ""
                        dotted_fields = self._detect_dotted_leader_fields(page_text, gray.shape, page_num)
                    all_fields.extend(dotted_fields)
                except Exception as _dot_err:
                    print(f"Error detecting dotted leaders on page {page_num}: {_dot_err}")
            
            doc.close()
            return all_fields
            
        except Exception as e:
            print(f"Error in OCR field detection: {e}")
            return []
    
    def _detect_form_fields_advanced(self, gray_image: np.ndarray, ocr_data: Dict, page_num: int = 0) -> List[FormField]:
        """Advanced form field detection using multiple techniques"""
        fields = []
        
        # Method 1: Detect rectangular form fields
        rectangular_fields = self._detect_rectangular_fields(gray_image, page_num)
        
        # Method 2: Detect fields based on OCR text patterns
        text_based_fields = self._detect_text_based_fields(ocr_data, gray_image.shape, page_num)
        
        # Method 3: Detect underlines and form lines
        line_based_fields = self._detect_line_based_fields(gray_image, page_num)

        # Method 4: Detect dotted lines (dot leaders) visually
        dotted_line_fields = self._detect_dotted_lines(gray_image, page_num)
        
        # Combine and deduplicate
        all_fields = rectangular_fields + text_based_fields + line_based_fields + dotted_line_fields
        fields = self._deduplicate_fields(all_fields)
        
        return fields

    def _detect_dotted_lines(self, gray_image: np.ndarray, page_num: int = 0) -> List[FormField]:
        """Detect dotted leader lines visually and turn them into input fields above the line.

        This connects dot patterns horizontally and then finds long thin components.
        """
        fields: List[FormField] = []
        try:
            # Threshold: dark dots become white (foreground)
            _, bin_inv = cv2.threshold(gray_image, 200, 255, cv2.THRESH_BINARY_INV)

            # Bridge gaps horizontally to connect dotted segments
            horiz_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (25, 1))
            connected = cv2.morphologyEx(bin_inv, cv2.MORPH_CLOSE, horiz_kernel, iterations=1)

            contours, _ = cv2.findContours(connected, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            img_h, img_w = gray_image.shape

            idx = 0
            for cnt in contours:
                x, y, w, h = cv2.boundingRect(cnt)
                if w < 100 or h > 12:
                    continue

                # Sample text region above to classify
                y0 = max(0, y - 35)
                roi = gray_image[y0:y, x:min(img_w, x + w)]
                text = ''
                if roi.size > 0:
                    try:
                        text = pytesseract.image_to_string(roi, config='--psm 7').strip()
                    except Exception:
                        text = ''
                text_lower = (text or '').lower()

                field_type = 'text'
                if 'day' in text_lower:
                    field_type = 'day'
                elif 'month' in text_lower:
                    field_type = 'month'
                elif 'year' in text_lower or '20' in text_lower:
                    field_type = 'year'
                elif any(k in text_lower for k in ['employer', 'employee', 'name']):
                    field_type = 'name'

                field = FormField(
                    id=f"dots_vis_p{page_num}_{idx}",
                    field_type=field_type,
                    x=x,
                    y=max(0, y - 24),
                    width=w,
                    height=25,
                    context=text if text else 'dotted_line',
                    confidence=0.75
                )
                field.page = page_num
                fields.append(field)
                idx += 1

            return fields
        except Exception as e:
            print(f"Error detecting dotted lines: {e}")
            return fields

    def _detect_dotted_leader_fields(self, page_text: str, image_shape: Tuple, page_num: int = 0) -> List[FormField]:
        """Detect dotted leader placeholders like sequences of '.' or '…' in text and
        estimate input rectangles at those positions. Coordinates are in 3x pixel space.
        """
        fields: List[FormField] = []
        try:
            if not page_text:
                return fields

            lines = page_text.split('\n')
            img_h = image_shape[0] if isinstance(image_shape, tuple) and len(image_shape) >= 1 else 2400
            base_y = 50
            line_height = 30

            dotted_pattern = re.compile(r'(?:\.{3,}|…{1,})')
            field_id_counter = 0

            for line_idx, raw_line in enumerate(lines):
                line = raw_line.rstrip()
                if not line or ('.' not in line and '…' not in line):
                    continue

                for match in dotted_pattern.finditer(line):
                    start_pos = match.start()
                    end_pos = match.end()
                    run_text = line[start_pos:end_pos]

                    # Context
                    context_before = line[max(0, start_pos-40):start_pos]
                    context_after = line[end_pos:min(len(line), end_pos+40)]
                    combined_context = (context_before + ' ' + context_after).lower()

                    # Type guess
                    if 'day' in combined_context and 'month' not in combined_context:
                        field_type = 'day'
                        est_width = 80
                    elif 'month' in combined_context:
                        field_type = 'month'
                        est_width = 140
                    elif 'year' in combined_context or '20' in combined_context:
                        field_type = 'year'
                        est_width = 100
                    elif any(k in combined_context for k in ['employer', 'employee', 'company']):
                        field_type = 'name'
                        est_width = 260
                    else:
                        field_type = 'text'
                        est_width = 200

                    # Coordinate estimate from character positions
                    x_est = 50 + max(0, start_pos) * 6
                    y_est = base_y + line_idx * line_height
                    width_est = max(80, min(est_width, 400))
                    height_est = 25
                    y_est = min(max(0, y_est), max(0, img_h - height_est))

                    field = FormField(
                        id=f"dots_field_p{page_num}_{field_id_counter}",
                        field_type=field_type,
                        x=x_est,
                        y=y_est,
                        width=width_est,
                        height=height_est,
                        context=f"{context_before} [{run_text}] {context_after}"[:160],
                        confidence=0.85
                    )
                    field.page = page_num
                    fields.append(field)
                    field_id_counter += 1

            return fields
        except Exception as e:
            print(f"Error in dotted leader detection: {e}")
            return fields

    def _detect_dotted_leader_fields_pdf(self, page, page_num: int, scale: float = 3.0) -> List[FormField]:
        """Detect dotted leader placeholders using PDF raw text positions for accurate boxes.

        We scan spans and search for sequences of '.' or '…', then map their bbox to a field rect
        with a small height and a reasonable width based on run length.
        Returned coordinates are scaled by 'scale' to match image-space used downstream.
        """
        fields: List[FormField] = []
        try:
            raw = page.get_text('rawdict')
            if not raw or 'blocks' not in raw:
                return fields

            pattern = re.compile(r'(\.{3,}|…{1,})')
            field_id = 0

            for block in raw.get('blocks', []):
                for line in block.get('lines', []):
                    # Build line string and map char offsets to approximate x positions
                    spans = line.get('spans', [])
                    if not spans:
                        continue
                    line_text = ''.join(s.get('text', '') for s in spans)
                    if not line_text or ('.' not in line_text and '…' not in line_text):
                        continue

                    # Approximate per-char width using span width/len
                    char_positions = []  # list of (char, x0, x1, y0, y1)
                    cursor_x = min((s['bbox'][0] for s in spans), default=0.0)
                    base_y0 = min((s['bbox'][1] for s in spans), default=0.0)
                    base_y1 = max((s['bbox'][3] for s in spans), default=0.0)
                    built = ''
                    for s in spans:
                        t = s.get('text', '') or ''
                        x0, y0, x1, y1 = s.get('bbox', (0, 0, 0, 0))
                        width = max(0.1, x1 - x0)
                        per_char = width / max(1, len(t))
                        # Emit per char positions
                        cx = x0
                        for ch in t:
                            char_positions.append((ch, cx, cx + per_char, y0, y1))
                            cx += per_char
                        built += t
                    if not built:
                        continue

                    # Search dotted runs within built
                    for m in pattern.finditer(built):
                        s_idx, e_idx = m.start(), m.end()
                        if s_idx >= len(char_positions):
                            continue
                        start_char = char_positions[s_idx]
                        end_char = char_positions[min(e_idx - 1, len(char_positions) - 1)]
                        x0 = start_char[1]
                        x1 = end_char[2]
                        y0 = min(start_char[3], end_char[3], base_y0)
                        y1 = max(start_char[4], end_char[4], base_y1)

                        # Expand to a usable field rectangle
                        est_w = max(80.0, min((x1 - x0) * 2.0, 400.0))
                        est_h = max(20.0, min((y1 - y0) * 1.2, 28.0))
                        rect_x = x0
                        rect_y = y0 - est_h  # place above the dots line a bit

                        # Context for classification
                        context_before = built[max(0, s_idx-40):s_idx]
                        context_after = built[e_idx:min(len(built), e_idx+40)]
                        combined = (context_before + ' ' + context_after).lower()
                        if 'day' in combined and 'month' not in combined:
                            ftype = 'day'
                        elif 'month' in combined:
                            ftype = 'month'
                        elif 'year' in combined or '20' in combined:
                            ftype = 'year'
                        elif any(k in combined for k in ['employer', 'employee', 'company']):
                            ftype = 'name'
                        else:
                            ftype = 'text'

                        # Scale to image coordinates (3x) and clamp to page bounds
                        page_w, page_h = float(page.rect.width), float(page.rect.height)
                        sx = max(0.0, min(rect_x, page_w)) * scale
                        sy = max(0.0, min(rect_y, page_h)) * scale
                        sw = max(40.0, min(est_w, page_w - rect_x)) * scale
                        sh = max(16.0, min(est_h, page_h * 0.1)) * scale

                        field = FormField(
                            id=f"dots_pdf_p{page_num}_{field_id}",
                            field_type=ftype,
                            x=int(sx),
                            y=int(sy),
                            width=int(sw),
                            height=int(sh),
                            context=f"{context_before} [{m.group(0)}] {context_after}"[:160],
                            confidence=0.9
                        )
                        field.page = page_num
                        fields.append(field)
                        field_id += 1

            return fields
        except Exception as e:
            print(f"Error in PDF dotted leader detection: {e}")
            return fields
    
    def _detect_rectangular_fields(self, gray_image: np.ndarray, page_num: int = 0) -> List[FormField]:
        """Detect rectangular form fields using contour analysis"""
        fields = []
        try:
            # Apply adaptive threshold
            thresh = cv2.adaptiveThreshold(gray_image, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                         cv2.THRESH_BINARY_INV, 11, 2)
            
            # Find contours
            contours, _ = cv2.findContours(thresh, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            image_height, image_width = gray_image.shape
            
            for i, contour in enumerate(contours):
                x, y, w, h = cv2.boundingRect(contour)
                area = w * h
                aspect_ratio = w / h if h > 0 else 0
                
                # Filter for form field characteristics
                if (1000 < area < 50000 and  # Reasonable size
                    30 < w < image_width * 0.7 and  # Width constraints
                    15 < h < image_height * 0.2 and  # Height constraints
                    0.3 < aspect_ratio < 15):  # Aspect ratio constraints
                    
                    # Check if area is mostly blank
                    roi = gray_image[y:y+h, x:x+w]
                    if roi.size > 0:
                        mean_intensity = np.mean(roi)
                        if mean_intensity > 200:  # Mostly white
                            field_type = self._classify_field_by_context(gray_image, x, y, w, h)
                            
                            field = FormField(
                                id=f"rect_field_p{page_num}_{i}",
                                field_type=field_type,
                                x=x,
                                y=y,
                                width=w,
                                height=h,
                                context=field_type,
                                confidence=0.7
                            )
                            # Store page number as a custom attribute
                            field.page = page_num
                            fields.append(field)
            
            return fields
            
        except Exception as e:
            print(f"Error detecting rectangular fields: {e}")
            return []
    
    def _detect_text_based_fields(self, ocr_data: Dict, image_shape: Tuple, page_num: int = 0) -> List[FormField]:
        """Detect fields based on OCR text patterns"""
        fields = []
        try:
            height, width = image_shape
            
            # Look for text patterns that indicate form fields
            field_indicators = {
                'name': ['name:', 'enter your name', 'full name', 'first name', 'last name'],
                'email': ['email:', 'e-mail:', 'email address'],
                'phone': ['phone:', 'telephone:', 'tel:', 'mobile:', 'number:'],
                'address': ['address:', 'street:', 'location:'],
                'date': ['date:', 'birth date', 'dob:', 'date of birth'],
                'age': ['age:', 'years old'],
                'signature': ['signature:', 'sign:', 'initial:']
            }
            
            # Get text with bounding boxes
            n_boxes = len(ocr_data['text'])
            for i in range(n_boxes):
                text = ocr_data['text'][i].strip().lower()
                conf = int(ocr_data['conf'][i])
                
                if conf > 30:  # Only process high-confidence text
                    for field_type, indicators in field_indicators.items():
                        if any(indicator in text for indicator in indicators):
                            # Look for blank space after this text
                            x = ocr_data['left'][i]
                            y = ocr_data['top'][i]
                            w = ocr_data['width'][i]
                            h = ocr_data['height'][i]
                            
                            # Estimate field position (usually to the right or below)
                            field_x = x + w + 10
                            field_y = y
                            field_w = 200  # Default width
                            field_h = max(h, 25)  # At least as tall as text
                            
                            # Ensure field is within image bounds
                            if field_x + field_w > width:
                                field_x = x
                                field_y = y + h + 5
                            
                            field = FormField(
                                id=f"text_field_p{page_num}_{i}",
                                field_type=field_type,
                                x=field_x,
                                y=field_y,
                                width=field_w,
                                height=field_h,
                                context=field_type,
                                confidence=0.8
                            )
                            # Store page number as a custom attribute
                            field.page = page_num
                            fields.append(field)
                            break
            
            return fields
            
        except Exception as e:
            print(f"Error detecting text-based fields: {e}")
            return []
    
    def _detect_line_based_fields(self, gray_image: np.ndarray, page_num: int = 0) -> List[FormField]:
        """Detect fields based on underlines and form lines"""
        fields = []
        try:
            # Detect horizontal lines
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            horizontal_lines = cv2.morphologyEx(gray_image, cv2.MORPH_OPEN, horizontal_kernel)
            
            # Find contours of horizontal lines
            contours, _ = cv2.findContours(horizontal_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            for i, contour in enumerate(contours):
                x, y, w, h = cv2.boundingRect(contour)
                
                # Look for text above the line
                text_region = gray_image[max(0, y-30):y, x:x+w]
                if text_region.size > 0:
                    # Use OCR to detect text above the line
                    text = pytesseract.image_to_string(text_region, config='--psm 8').strip()
                    
                    if text:
                        field_type = self._classify_text_to_field_type(text)
                        
                        field = FormField(
                            id=f"line_field_p{page_num}_{i}",
                            field_type=field_type,
                            x=x,
                            y=y-25,
                            width=w,
                            height=25,
                            context=field_type,
                            confidence=0.6
                        )
                        # Store page number as a custom attribute
                        field.page = page_num
                        fields.append(field)
            
            return fields
            
        except Exception as e:
            print(f"Error detecting line-based fields: {e}")
            return []
    
    def _classify_field_by_context(self, gray_image: np.ndarray, x: int, y: int, w: int, h: int) -> str:
        """Classify field type based on surrounding context"""
        try:
            # Extract text around the field
            context_region = gray_image[max(0, y-20):min(gray_image.shape[0], y+h+20), 
                                      max(0, x-50):min(gray_image.shape[1], x+w+50)]
            
            if context_region.size > 0:
                context_text = pytesseract.image_to_string(context_region).lower()
                
                # Classify based on context
                if any(keyword in context_text for keyword in ['name', 'enter your name']):
                    return 'name'
                elif any(keyword in context_text for keyword in ['email', 'e-mail']):
                    return 'email'
                elif any(keyword in context_text for keyword in ['phone', 'telephone', 'tel']):
                    return 'phone'
                elif any(keyword in context_text for keyword in ['address', 'street']):
                    return 'address'
                elif any(keyword in context_text for keyword in ['date', 'birth', 'dob']):
                    return 'date'
                elif any(keyword in context_text for keyword in ['age', 'years']):
                    return 'age'
                elif any(keyword in context_text for keyword in ['signature', 'sign']):
                    return 'signature'
            
            return 'text'
            
        except Exception as e:
            print(f"Error classifying field context: {e}")
            return 'text'
    
    def _classify_text_to_field_type(self, text: str) -> str:
        """Classify field type based on text content"""
        text_lower = text.lower().strip()
        
        if any(keyword in text_lower for keyword in ['name', 'enter your name', 'full name']):
            return 'name'
        elif any(keyword in text_lower for keyword in ['email', 'e-mail']):
            return 'email'
        elif any(keyword in text_lower for keyword in ['phone', 'telephone', 'tel', 'mobile']):
            return 'phone'
        elif any(keyword in text_lower for keyword in ['address', 'street']):
            return 'address'
        elif any(keyword in text_lower for keyword in ['date', 'birth', 'dob']):
            return 'date'
        elif any(keyword in text_lower for keyword in ['age', 'years old']):
            return 'age'
        elif any(keyword in text_lower for keyword in ['signature', 'sign']):
            return 'signature'
        else:
            return 'text'
    
    def _deduplicate_fields(self, fields: List[FormField]) -> List[FormField]:
        """Remove duplicate fields based on position overlap"""
        if not fields:
            return []
        
        # Sort by confidence (highest first)
        fields.sort(key=lambda f: f.confidence, reverse=True)
        
        unique_fields = []
        for field in fields:
            is_duplicate = False
            for existing_field in unique_fields:
                # Check for overlap
                if self._fields_overlap(field, existing_field):
                    is_duplicate = True
                    break
            
            if not is_duplicate:
                unique_fields.append(field)
        
        return unique_fields
    
    def _fields_overlap(self, field1: FormField, field2: FormField, threshold: float = 0.3) -> bool:
        """Check if two fields overlap significantly"""
        # Calculate intersection area
        x1 = max(field1.x, field2.x)
        y1 = max(field1.y, field2.y)
        x2 = min(field1.x + field1.width, field2.x + field2.width)
        y2 = min(field1.y + field1.height, field2.y + field2.height)
        
        if x1 < x2 and y1 < y2:
            intersection_area = (x2 - x1) * (y2 - y1)
            field1_area = field1.width * field1.height
            field2_area = field2.width * field2.height
            
            # Calculate overlap ratio
            overlap_ratio = intersection_area / min(field1_area, field2_area)
            return overlap_ratio > threshold
        
        return False
    
    def _merge_field_detections(self, acroform_fields: List[FormField], ocr_fields: List[FormField]) -> List[FormField]:
        """Merge AcroForm and OCR field detections"""
        # Prioritize AcroForm fields (they're more accurate)
        all_fields = acroform_fields + ocr_fields
        return self._deduplicate_fields(all_fields)
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    def _process_image(self, file_path: str) -> Dict:
        """Process image file"""
        try:
            image = cv2.imread(file_path)
            if image is None:
                raise ValueError(f"Could not load image: {file_path}")
            
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            self.extracted_text = pytesseract.image_to_string(gray)
            
            # Detect fields using OCR
            ocr_data = pytesseract.image_to_data(gray, output_type=pytesseract.Output.DICT)
            fields = self._detect_form_fields_advanced(gray, ocr_data)
            
            return {
                'extracted_text': self.extracted_text,
                'fields': fields,
                'total_fields': len(fields),
                'document_type': 'image'
            }
            
        except Exception as e:
            raise Exception(f"Error processing image: {str(e)}")
    
    def _process_word(self, file_path: str) -> Dict:
        """Process Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            self.extracted_text = '\n'.join(text_content)
            
            # Detect form fields in Word document
            fields = self._detect_word_form_fields(doc)
            
            return {
                'extracted_text': self.extracted_text,
                'fields': fields,
                'total_fields': len(fields),
                'document_type': 'word'
            }
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def _detect_word_form_fields(self, doc) -> List[FormField]:
        """Detect form fields in Word document"""
        fields = []
        field_id = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Look for form field indicators
            field_type = self._classify_text_to_field_type(text)
            if field_type != 'text' or ':' in text:
                # This might be a form field
                field = FormField(
                    id=f"word_field_{field_id}",
                    field_type=field_type,
                    x=0,  # Word doesn't have precise positioning
                    y=field_id * 30,
                    width=200,
                    height=25,
                    context=field_type,
                    confidence=0.6
                )
                fields.append(field)
                field_id += 1
        
        return fields

def convert_form_fields_to_dict(fields: List[FormField]) -> List[Dict]:
    """Convert FormField objects to dictionary format for API compatibility"""
    return [
        {
            'id': field.id,
            'field_type': field.field_type,
            'x_position': field.x,
            'y_position': field.y,
            'width': field.width,
            'height': field.height,
            'context': field.context,
            'confidence': field.confidence,
            'is_required': field.is_required,
            'placeholder': field.placeholder,
            'validation_pattern': field.validation_pattern,
            'page': getattr(field, 'page', 0),  # Include page number
            'user_content': '',
            'ai_suggestion': '',
            'ai_enhanced': False
        }
        for field in fields
    ]
